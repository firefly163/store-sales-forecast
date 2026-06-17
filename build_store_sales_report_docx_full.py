from __future__ import annotations

from pathlib import Path
import math

import pandas as pd
from PIL import Image, ImageDraw

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from build_store_sales_report_docx import (
    ASSET_DIR,
    TABLE_DIR,
    WORD_FIG_DIR,
    OUT_DOCX,
    configure_doc,
    read_csv,
    get_font,
    draw_title,
    text_size,
    nice_max,
    add_heading,
    add_para,
    add_bullet,
    add_picture,
    add_df_table,
    set_run_font,
    line_chart_monthly,
    bar_chart_top_family,
    bar_chart_store_type,
    score_progression_chart,
    component_weights_chart,
    workflow_chart,
    submission_distribution_chart,
)


OUT_FULL_DOCX = ASSET_DIR / "Store_Sales_实验报告_完整版.docx"

BLUE = (31, 78, 121)
ORANGE = (237, 125, 49)
GREEN = (112, 173, 71)
PURPLE = (112, 48, 160)
GRAY = (90, 90, 90)
LIGHT_GRAY = (245, 247, 250)


def draw_wrapped_center(draw, center_x, y, text, font, fill, line_height=28, max_chars=15):
    lines = []
    for raw in str(text).split("\n"):
        while len(raw) > max_chars:
            lines.append(raw[:max_chars])
            raw = raw[max_chars:]
        lines.append(raw)
    for i, line in enumerate(lines):
        tw, _ = text_size(draw, line, font)
        draw.text((center_x - tw / 2, y + i * line_height), line, fill=fill, font=font)


def data_relationship_chart() -> Path:
    out = WORD_FIG_DIR / "data_relationship.png"
    w, h = 1300, 760
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "数据表关系与建模输入", w)

    font = get_font(21, True)
    small = get_font(18)
    boxes = [
        ("train.csv\n历史销量", 60, 120, BLUE),
        ("test.csv\n待预测id", 60, 260, BLUE),
        ("stores.csv\n门店属性", 60, 400, ORANGE),
        ("oil.csv\n每日油价", 60, 540, ORANGE),
        ("holidays_events.csv\n节假日事件", 365, 120, GREEN),
        ("transactions.csv\n客流交易量", 365, 260, GREEN),
        ("日期特征\n周/月/年/工资日", 365, 400, PURPLE),
        ("特征矩阵 data\nstore × family × date", 690, 255, BLUE),
        ("模型训练与验证\nLGBM/XGB/CatBoost", 995, 205, ORANGE),
        ("submission.csv\nid, sales", 995, 430, GREEN),
    ]
    for text, x, y, color in boxes:
        draw.rounded_rectangle((x, y, x + 230, y + 90), radius=16, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 115, y + 18, text, font, color, line_height=30, max_chars=16)

    arrows = [
        (290, 165, 690, 300), (290, 305, 690, 300), (290, 445, 690, 300),
        (290, 585, 690, 300), (595, 165, 690, 300), (595, 305, 690, 300),
        (595, 445, 690, 300), (920, 300, 995, 250), (1110, 295, 1110, 430),
    ]
    for x1, y1, x2, y2 in arrows:
        draw.line((x1, y1, x2, y2), fill=(120, 120, 120), width=3)
        draw.ellipse((x2 - 4, y2 - 4, x2 + 4, y2 + 4), fill=(120, 120, 120))

    draw.text((80, 680), "核心样本粒度：一个日期 date × 一个门店 store_nbr × 一个商品品类 family，对应一个 sales 预测值。", fill=GRAY, font=small)
    img.save(out, quality=95)
    return out


def validation_timeline_chart() -> Path:
    out = WORD_FIG_DIR / "validation_timeline.png"
    w, h = 1200, 420
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "时间序列训练、验证与测试切分", w)

    y = 210
    x0, x1, x2, x3 = 95, 810, 970, 1120
    draw.line((x0, y, x3, y), fill=(180, 180, 180), width=6)
    draw.rounded_rectangle((x0, y - 32, x1, y + 32), radius=14, fill=(221, 235, 247), outline=BLUE, width=3)
    draw.rounded_rectangle((x1, y - 32, x2, y + 32), radius=14, fill=(252, 228, 214), outline=ORANGE, width=3)
    draw.rounded_rectangle((x2, y - 32, x3, y + 32), radius=14, fill=(226, 239, 218), outline=GREEN, width=3)

    font = get_font(21, True)
    small = get_font(18)
    draw_wrapped_center(draw, (x0 + x1) / 2, y - 12, "训练区间\n2013-01-01 至 2017-07-15", font, BLUE, 28, 30)
    draw_wrapped_center(draw, (x1 + x2) / 2, y - 12, "验证窗口\n31天", font, ORANGE, 28, 20)
    draw_wrapped_center(draw, (x2 + x3) / 2, y - 12, "测试窗口\n16天", font, GREEN, 28, 20)
    draw.text((95, 315), "为什么不能随机切分：随机切分会把未来信息泄漏到训练集，本题必须按时间顺序验证模型的外推能力。", fill=GRAY, font=small)
    img.save(out, quality=95)
    return out


def rmsle_explanation_chart() -> Path:
    out = WORD_FIG_DIR / "rmsle_explanation.png"
    w, h = 1200, 560
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "RMSLE 指标与建模启发", w)
    font = get_font(24, True)
    small = get_font(19)

    formula = "RMSLE = sqrt(mean((log(1 + y_pred) - log(1 + y_true))^2))"
    tw, _ = text_size(draw, formula, font)
    draw.rounded_rectangle((90, 115, 1110, 185), radius=18, fill=LIGHT_GRAY, outline=BLUE, width=3)
    draw.text(((w - tw) / 2, 135), formula, fill=BLUE, font=font)

    notes = [
        ("对 sales 取 log1p", "模型训练目标使用 log1p(sales)，与评价指标保持一致。", BLUE),
        ("关注相对误差", "预测 10 成 20 与预测 1000 成 2000 都会被明显惩罚。", ORANGE),
        ("天然要求非负", "提交前必须 clip 到非负，并处理闭店、长期零销量组合。", GREEN),
        ("极端值影响降低", "异常大销量仍要处理，但不会像 RMSE 那样完全支配训练。", PURPLE),
    ]
    for i, (title, body, color) in enumerate(notes):
        x = 90 + i * 270
        draw.rounded_rectangle((x, 250, x + 240, 430), radius=18, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 120, 275, title, font, color, 30, 12)
        draw_wrapped_center(draw, x + 120, 335, body, small, GRAY, 27, 12)
    img.save(out, quality=95)
    return out


def v7_feature_chart() -> Path:
    out = WORD_FIG_DIR / "v7_feature_engineering.png"
    w, h = 1350, 820
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V7 自研特征工程结构", w)
    title_font = get_font(22, True)
    body_font = get_font(17)

    groups = [
        ("基础日历", "day/week/month\ndayofweek/weekend\n工资日/月末", BLUE),
        ("外部事件", "油价插值\n节假日类型\n节前节后距离", ORANGE),
        ("时序滞后", "lag 16/21/28/56/112/364\nrolling mean/std\nEWM/momentum", GREEN),
        ("周期趋势", "Fourier 3.5/7/30/365\nTrend/Trend²\nFamily×Trend", PURPLE),
        ("层级交互", "store-family统计\ncity/state/cluster交互\nfamily/store rank", (47, 117, 181)),
        ("促销增强", "onpromotion lag\npromo rolling\npromo × holiday", (112, 173, 71)),
        ("异常与零销量", "99.5% clip\nclosed-store检测\n长期零销量置零", (192, 80, 77)),
        ("模型输入", "log1p(sales)\n类别编码\nfloat32压缩", (91, 155, 213)),
    ]
    for i, (title, body, color) in enumerate(groups):
        row, col = divmod(i, 4)
        x = 55 + col * 320
        y = 140 + row * 245
        draw.rounded_rectangle((x, y, x + 270, y + 180), radius=18, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 135, y + 22, title, title_font, color, 30, 12)
        draw_wrapped_center(draw, x + 135, y + 78, body, body_font, GRAY, 27, 16)

    draw.rounded_rectangle((170, 655, 1180, 725), radius=18, fill=(221, 235, 247), outline=BLUE, width=3)
    draw_wrapped_center(draw, 675, 675, "目标：把原始表格转换成能够表达趋势、周期、促销、节假日冲击和门店-品类差异的监督学习特征矩阵", body_font, BLUE, 28, 45)
    img.save(out, quality=95)
    return out


def ensemble_pipeline_chart() -> Path:
    out = WORD_FIG_DIR / "ensemble_pipeline.png"
    w, h = 1300, 650
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "模型训练、融合与后处理流程", w)
    font = get_font(22, True)
    small = get_font(18)

    boxes = [
        ("特征矩阵\nX_train/X_val/X_test", 65, 150, BLUE),
        ("LightGBM\nOptuna + 3 seeds", 355, 85, ORANGE),
        ("XGBoost\n多组深度/学习率", 355, 245, ORANGE),
        ("CatBoost\n类别特征补充", 355, 405, ORANGE),
        ("log-space 融合\nΣ wi·log_pred_i", 690, 245, GREEN),
        ("后处理\n非负截断/闭店置零\n季节性校正", 960, 210, PURPLE),
        ("最终提交\nid, sales", 960, 430, GREEN),
    ]
    for text, x, y, color in boxes:
        draw.rounded_rectangle((x, y, x + 240, y + 105), radius=18, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 120, y + 22, text, font if "\n" not in text[:8] else small, color if x != 65 else BLUE, 28, 18)

    for y in [137, 297, 457]:
        draw.line((305, 202, 355, y), fill=(120, 120, 120), width=3)
        draw.line((595, y, 690, 297), fill=(120, 120, 120), width=3)
    draw.line((930, 297, 960, 262), fill=(120, 120, 120), width=3)
    draw.line((1080, 315, 1080, 430), fill=(120, 120, 120), width=3)
    draw.text((80, 570), "验证集上先比较单模型，再搜索融合权重；提交前检查预测分布，避免负值、异常全零或极端漂移。", fill=GRAY, font=small)
    img.save(out, quality=95)
    return out


def promo_oil_chart() -> Path:
    df = read_csv("monthly_sales_promo_oil.csv")
    out = WORD_FIG_DIR / "promo_oil_sales.png"
    w, h = 1250, 650
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "销售额、促销数量与油价的月度关系", w)

    left, right, top, bottom = 85, 70, 95, 115
    plot_w, plot_h = w - left - right, h - top - bottom
    sales = df["sales"].astype(float).values
    oil = df["oil"].astype(float).values
    promo = df["onpromotion"].astype(float).values
    n = len(df)

    draw.rectangle((left, top, left + plot_w, top + plot_h), outline=(210, 210, 210), width=2)
    for i in range(5):
        y = top + i * plot_h / 4
        draw.line((left, y, left + plot_w, y), fill=(232, 232, 232), width=1)

    def pts(values, invert=False):
        mn, mx = float(values.min()), float(values.max())
        if mx == mn:
            mx = mn + 1
        p = []
        for i, v in enumerate(values):
            x = left + i * plot_w / max(1, n - 1)
            y = top + plot_h - (float(v) - mn) / (mx - mn) * plot_h
            p.append((x, y))
        return p

    draw.line(pts(sales), fill=BLUE, width=4)
    draw.line(pts(oil), fill=ORANGE, width=3)
    draw.line(pts(promo), fill=GREEN, width=3)

    font = get_font(18)
    draw.rectangle((870, 110, 1180, 190), fill="white", outline=(210, 210, 210))
    for j, (name, color) in enumerate([("sales", BLUE), ("oil", ORANGE), ("onpromotion", GREEN)]):
        y = 124 + j * 23
        draw.line((890, y + 8, 930, y + 8), fill=color, width=4)
        draw.text((940, y), name, fill=GRAY, font=font)

    for i in range(0, n, max(1, n // 8)):
        x = left + i * plot_w / max(1, n - 1)
        label = str(df["month"].iloc[i])[:7]
        tw, _ = text_size(draw, label, font)
        draw.text((x - tw / 2, top + plot_h + 18), label, fill=GRAY, font=font)

    draw.text((left, h - 42), "三条曲线分别做归一化展示，用于说明外部变量和销售波动之间存在共同变化，需要进入模型。", fill=GRAY, font=font)
    img.save(out, quality=95)
    return out


def holiday_bar_chart() -> Path:
    df = read_csv("holiday_type_counts.csv")
    out = WORD_FIG_DIR / "holiday_type_counts.png"
    w, h = 1050, 560
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "节假日事件类型分布", w)
    left, base_y = 110, 445
    plot_w, plot_h = 820, 300
    max_v = nice_max(float(df["count"].max()))
    font = get_font(19)
    colors = [BLUE, ORANGE, GREEN, PURPLE, (91, 155, 213), (192, 80, 77)]
    for i, row in df.iterrows():
        x = left + i * plot_w / len(df) + 25
        bw = plot_w / len(df) - 50
        val = float(row["count"])
        bh = plot_h * val / max_v
        draw.rounded_rectangle((x, base_y - bh, x + bw, base_y), radius=8, fill=colors[i % len(colors)])
        draw_wrapped_center(draw, x + bw / 2, base_y + 12, row["holiday_type"], font, GRAY, 22, 10)
        draw_wrapped_center(draw, x + bw / 2, base_y - bh - 27, int(val), font, GRAY, 22, 10)
    img.save(out, quality=95)
    return out


def cluster_bar_chart() -> Path:
    df = read_csv("store_cluster_counts.csv")
    out = WORD_FIG_DIR / "store_cluster_counts.png"
    w, h = 1200, 560
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "门店 Cluster 分布", w)
    left, base_y = 80, 450
    plot_w, plot_h = 1060, 300
    max_v = nice_max(float(df["store_count"].max()))
    font = get_font(16)
    for i, row in df.iterrows():
        x = left + i * plot_w / len(df) + 8
        bw = plot_w / len(df) - 16
        val = float(row["store_count"])
        bh = plot_h * val / max_v
        color = BLUE if i % 2 == 0 else ORANGE
        draw.rounded_rectangle((x, base_y - bh, x + bw, base_y), radius=6, fill=color)
        label = str(row["cluster"])
        tw, _ = text_size(draw, label, font)
        draw.text((x + bw / 2 - tw / 2, base_y + 10), label, fill=GRAY, font=font)
        count = str(int(val))
        tw, _ = text_size(draw, count, font)
        draw.text((x + bw / 2 - tw / 2, base_y - bh - 22), count, fill=GRAY, font=font)
    draw.text((left, h - 35), "cluster 是门店层级的重要结构变量，V7 中使用 cluster_rank 与 cluster × family / cluster × trend 交互。", fill=GRAY, font=get_font(18))
    img.save(out, quality=95)
    return out


def v1_baseline_pipeline_chart() -> Path:
    out = WORD_FIG_DIR / "v1_baseline_pipeline.png"
    w, h = 1300, 620
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V1 基础版建模流程", w)
    title_font = get_font(22, True)
    body_font = get_font(18)
    boxes = [
        ("读取数据", "train/test\nstores/oil/holiday\ntransactions", 55, 170, BLUE),
        ("基础清洗", "日期解析\n字符串标准化\n缺失油价填充", 315, 170, ORANGE),
        ("基础特征", "门店属性\n节假日透视\n简单 lag", 575, 170, GREEN),
        ("基线模型", "XGBoost\nLightGBM\n线性/GBDT对照", 835, 170, PURPLE),
        ("提交文件", "submission.csv\n查看分数\n记录问题", 1095, 170, BLUE),
    ]
    for title, body, x, y, color in boxes:
        draw.rounded_rectangle((x, y, x + 190, y + 205), radius=18, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 95, y + 24, title, title_font, color, 30, 10)
        draw_wrapped_center(draw, x + 95, y + 84, body, body_font, GRAY, 28, 18)
    for x in [250, 510, 770, 1030]:
        draw.line((x, 272, x + 55, 272), fill=(120, 120, 120), width=4)
        draw.polygon([(x + 55, 272), (x + 38, 262), (x + 38, 282)], fill=(120, 120, 120))
    draw.text((70, 500), "V1 的定位：先跑通竞赛闭环，确认数据合并、特征生成、模型训练和提交格式没有问题。", fill=GRAY, font=get_font(20))
    img.save(out, quality=95)
    return out


def v1_feature_scope_chart() -> Path:
    out = WORD_FIG_DIR / "v1_feature_scope.png"
    w, h = 1200, 620
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V1 基础特征范围", w)
    title_font = get_font(22, True)
    body_font = get_font(18)
    items = [
        ("日期特征", "year / month / day\n星期信息\n基础时间字段", BLUE),
        ("门店特征", "city / state / type\ncluster\nstore_nbr", ORANGE),
        ("外部变量", "油价 ffill/bfill\n节假日透视\ntransactions", GREEN),
        ("基础时序", "少量 lag\n简单历史均值\n未来 lag 补齐", PURPLE),
    ]
    for i, (title, body, color) in enumerate(items):
        x = 120 + i * 260
        draw.rounded_rectangle((x, 150, x + 210, 360), radius=20, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 105, 178, title, title_font, color, 30, 12)
        draw_wrapped_center(draw, x + 105, 245, body, body_font, GRAY, 28, 14)
    draw.rounded_rectangle((180, 430, 1020, 500), radius=18, fill=(221, 235, 247), outline=BLUE, width=3)
    draw_wrapped_center(draw, 600, 448, "结论：V1 能形成 baseline，但对强季节性、节前节后效应、门店-品类差异刻画不足。", body_font, BLUE, 28, 42)
    img.save(out, quality=95)
    return out


def v1_problem_diagnosis_chart() -> Path:
    out = WORD_FIG_DIR / "v1_problem_diagnosis.png"
    w, h = 1200, 620
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V1 结果诊断：为什么需要 V2", w)
    title_font = get_font(22, True)
    body_font = get_font(18)
    cards = [
        ("周期不足", "周/月/年周期没有被充分表达，节假日前后波动拟合弱。", BLUE),
        ("层级不足", "store-family、city-family、cluster-family 差异没有系统编码。", ORANGE),
        ("验证不足", "随机或短窗口验证容易高估效果，测试期外推不稳定。", GREEN),
        ("后处理不足", "闭店、长期零销量和极端销量处理不够细。", PURPLE),
    ]
    for i, (title, body, color) in enumerate(cards):
        x = 90 + i * 275
        draw.rounded_rectangle((x, 165, x + 235, 390), radius=18, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 117, 195, title, title_font, color, 30, 12)
        draw_wrapped_center(draw, x + 117, 265, body, body_font, GRAY, 28, 13)
    draw.text((100, 490), "改进方向：在 V2 中围绕周期、趋势、层级交互、时序滞后和后处理进行系统增强。", fill=GRAY, font=get_font(20))
    img.save(out, quality=95)
    return out


def v2_upgrade_map_chart() -> Path:
    out = WORD_FIG_DIR / "v2_upgrade_map.png"
    w, h = 1350, 720
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V2 相比 V1 的核心升级", w)
    title_font = get_font(21, True)
    body_font = get_font(18)
    rows = [
        ("时间表达", "基础 date 字段", "Fourier/趋势\n工资日/月末", BLUE),
        ("历史信息", "少量 lag", "多周期 lag\nrolling/EWM", ORANGE),
        ("层级结构", "门店/品类类别", "store-family\ncluster-family\ncity-family 统计", GREEN),
        ("异常处理", "简单缺失填充", "99.5% clip\nis_clipped 标记", PURPLE),
        ("模型策略", "单模型或简单对比", "LGB 多 seed\nXGB 多配置\nCatBoost 融合", (47, 117, 181)),
    ]
    draw.text((80, 120), "V1", fill=GRAY, font=title_font)
    draw.text((595, 120), "升级方向", fill=GRAY, font=title_font)
    draw.text((1030, 120), "V2/V7", fill=GRAY, font=title_font)
    for i, (topic, old, new, color) in enumerate(rows):
        y = 170 + i * 95
        draw.rounded_rectangle((60, y, 330, y + 78), radius=14, fill=LIGHT_GRAY, outline=color, width=2)
        draw_wrapped_center(draw, 195, y + 18, old, body_font, GRAY, 22, 18)
        draw.line((340, y + 39, 980, y + 39), fill=color, width=4)
        draw.polygon([(980, y + 39), (960, y + 28), (960, y + 50)], fill=color)
        draw_wrapped_center(draw, 660, y + 8, topic, title_font, color, 24, 12)
        draw.rounded_rectangle((995, y, 1285, y + 78), radius=14, fill=LIGHT_GRAY, outline=color, width=2)
        draw_wrapped_center(draw, 1140, y + 9, new, body_font, GRAY, 21, 22)
    img.save(out, quality=95)
    return out


def v2_lag_leakage_chart() -> Path:
    out = WORD_FIG_DIR / "v2_lag_no_leakage.png"
    w, h = 1300, 620
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V2 lag 设计：避免未来信息泄漏", w)
    font = get_font(20, True)
    small = get_font(17)
    x0, y = 100, 285
    step = 34
    for i in range(45):
        x = x0 + i * step
        day = i - 28
        if day < 0:
            color = (221, 235, 247)
            outline = BLUE
        elif 0 <= day < 16:
            color = (226, 239, 218)
            outline = GREEN
        else:
            color = (245, 245, 245)
            outline = (200, 200, 200)
        draw.rectangle((x, y, x + 26, y + 46), fill=color, outline=outline)
        if day in [-28, -21, -16, -7, 0, 15]:
            draw.text((x - 5, y + 55), str(day), fill=GRAY, font=small)
    draw.text((100, 190), "历史区间", fill=BLUE, font=font)
    draw.text((100 + 28 * step, 190), "预测窗口 16 天", fill=GREEN, font=font)
    draw.line((100 + 28 * step, 240, 100 + 28 * step, 370), fill=ORANGE, width=4)
    draw.text((100 + 28 * step - 75, 390), "预测起点", fill=ORANGE, font=font)
    allowed = "可用 lag: 16,17,18,19,20,21,28,35,56,84,112,182,364"
    draw.rounded_rectangle((120, 455, 1180, 525), radius=18, fill=LIGHT_GRAY, outline=BLUE, width=3)
    draw_wrapped_center(draw, 650, 475, allowed, font, BLUE, 28, 55)
    img.save(out, quality=95)
    return out


def v2_postprocess_chart() -> Path:
    out = WORD_FIG_DIR / "v2_postprocess.png"
    w, h = 1250, 620
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V2 后处理策略", w)
    title_font = get_font(22, True)
    body_font = get_font(18)
    boxes = [
        ("模型预测", "log-space ensemble\n还原到 sales", 80, 210, BLUE),
        ("非负约束", "clip(0, None)\n符合销量含义", 340, 210, ORANGE),
        ("闭店/零销量", "长期零销量组合\nclosed-store 置零", 600, 210, GREEN),
        ("季节性校正", "历史同期 lag\nfamily 自适应融合", 860, 210, PURPLE),
    ]
    for title, body, x, y, color in boxes:
        draw.rounded_rectangle((x, y, x + 210, y + 180), radius=18, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 105, y + 25, title, title_font, color, 30, 12)
        draw_wrapped_center(draw, x + 105, y + 92, body, body_font, GRAY, 28, 14)
    for x in [295, 555, 815]:
        draw.line((x, 300, x + 35, 300), fill=(120, 120, 120), width=4)
        draw.polygon([(x + 35, 300), (x + 20, 290), (x + 20, 310)], fill=(120, 120, 120))
    draw.text((90, 500), "后处理的目的不是“玄学调分”，而是把模型输出拉回业务约束：销量不能为负，闭店不应预测大量销售。", fill=GRAY, font=get_font(19))
    img.save(out, quality=95)
    return out


def v3_component_architecture_chart() -> Path:
    out = WORD_FIG_DIR / "v3_component_architecture.png"
    w, h = 1300, 650
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V3 组件化融合结构", w)
    title_font = get_font(21, True)
    body_font = get_font(18)
    comps = [
        ("LGB full", "全时段训练\n稳定主组件", BLUE),
        ("LGB 2015+", "近年窗口\n贴近测试分布", ORANGE),
        ("XGB full", "差异模型\n补充非线性", GREEN),
        ("XGB 2015+", "候选组件\n权重搜索检验", PURPLE),
    ]
    for i, (title, body, color) in enumerate(comps):
        x = 80 + i * 285
        draw.rounded_rectangle((x, 150, x + 220, 170 + 180), radius=18, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 110, 178, title, title_font, color, 30, 12)
        draw_wrapped_center(draw, x + 110, 245, body, body_font, GRAY, 28, 14)
        draw.line((x + 110, 350, 650, 435), fill=(120, 120, 120), width=3)
    draw.rounded_rectangle((430, 430, 870, 525), radius=20, fill=(226, 239, 218), outline=GREEN, width=3)
    draw_wrapped_center(draw, 650, 455, "log-space 加权融合\npred = expm1(Σ wi · log1p(pi))", title_font, GREEN, 30, 42)
    draw.text((90, 585), "V3 的关键不是盲目增加模型，而是把候选结果拆成组件，再用统一指标和分布检查选择权重。", fill=GRAY, font=get_font(19))
    img.save(out, quality=95)
    return out


def v3_weight_search_process_chart() -> Path:
    out = WORD_FIG_DIR / "v3_weight_search_process.png"
    df = read_csv("best_component_weights.csv")
    w, h = 1200, 620
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V3 权重搜索过程示意", w)
    font = get_font(21, True)
    small = get_font(18)
    steps = [
        ("候选提交", "读取各组件 sales\n按 id 对齐", BLUE),
        ("搜索空间", "权重和 = 1\n步长网格搜索", ORANGE),
        ("融合方式", "linear vs log\n优先 log-space", GREEN),
        ("筛选标准", "本地 RMSLE\n分布稳定性", PURPLE),
    ]
    for i, (title, body, color) in enumerate(steps):
        x = 80 + i * 270
        draw.rounded_rectangle((x, 135, x + 220, 280), radius=18, fill=LIGHT_GRAY, outline=color, width=3)
        draw_wrapped_center(draw, x + 110, 158, title, font, color, 28, 12)
        draw_wrapped_center(draw, x + 110, 215, body, small, GRAY, 26, 12)
    for x in [305, 575, 845]:
        draw.line((x, 207, x + 35, 207), fill=(120, 120, 120), width=4)
        draw.polygon([(x + 35, 207), (x + 20, 197), (x + 20, 217)], fill=(120, 120, 120))
    left, base_y = 180, 520
    plot_w, plot_h = 820, 170
    colors = [BLUE, ORANGE, GREEN, PURPLE]
    for i, row in df.iterrows():
        x = left + i * plot_w / len(df) + 55
        bw = plot_w / len(df) - 110
        val = float(row["weight"])
        bh = plot_h * val / 0.4
        draw.rounded_rectangle((x, base_y - bh, x + bw, base_y), radius=8, fill=colors[i])
        draw_wrapped_center(draw, x + bw / 2, base_y + 12, row["component"], small, GRAY, 24, 10)
        draw_wrapped_center(draw, x + bw / 2, base_y - bh - 26, f"{val:.1%}", small, GRAY, 24, 10)
    img.save(out, quality=95)
    return out


def v3_score_waterfall_chart() -> Path:
    out = WORD_FIG_DIR / "v3_score_waterfall.png"
    df = read_csv("score_progression.csv")
    labels = ["V2", "V3-1", "V3-2", "V3-3", "V3-4", "V3-5", "V3-6"]
    vals = df["public_score"].astype(float).values
    w, h = 1250, 650
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "V3 成绩推进瀑布图", w)
    left, top, plot_w, plot_h = 95, 120, 1080, 380
    ymin, ymax = 0.3788, 0.4010
    font = get_font(18)
    draw.rectangle((left, top, left + plot_w, top + plot_h), outline=(210, 210, 210), width=2)
    for i in range(5):
        y = top + i * plot_h / 4
        score = ymax - i * (ymax - ymin) / 4
        draw.line((left, y, left + plot_w, y), fill=(232, 232, 232), width=1)
        draw.text((25, y - 10), f"{score:.3f}", fill=GRAY, font=font)
    bar_w = plot_w / len(vals) * 0.55
    for i, v in enumerate(vals):
        x = left + i * plot_w / len(vals) + plot_w / len(vals) * 0.22
        y = top + (ymax - v) / (ymax - ymin) * plot_h
        color = GREEN if i == len(vals) - 1 else ORANGE
        draw.rounded_rectangle((x, y, x + bar_w, top + plot_h), radius=8, fill=color)
        draw_wrapped_center(draw, x + bar_w / 2, y - 28, f"{v:.5f}", font, GRAY, 22, 10)
        draw_wrapped_center(draw, x + bar_w / 2, top + plot_h + 18, labels[i], font, GRAY, 22, 10)
    draw.text((left, h - 65), "纵轴为 Public RMSLE，柱越高表示分数越低、效果越好；V3 后期提升虽小，但在榜单前段仍有意义。", fill=GRAY, font=font)
    img.save(out, quality=95)
    return out


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(45, 45, 45)
    doc.add_paragraph()


def add_simple_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Store Sales 时间序列预测实验报告")
    set_run_font(run, 24, True, BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Kaggle Store Sales - Time Series Forecasting | V1/V2/V3 迭代建模与提交优化")
    set_run_font(run, 12, False, GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("最终 Public RMSLE：0.37927    最高截图排名：约 298 名")
    set_run_font(run, 11, True, ORANGE)


def add_metadata(doc: Document) -> None:
    meta = pd.DataFrame(
        [
            ["竞赛题目", "Store Sales - Time Series Forecasting"],
            ["竞赛平台", "Kaggle"],
            ["赛题任务", "预测厄瓜多尔 Favorita 超市未来 16 天各门店-品类销量"],
            ["训练时间范围", "2013-01-01 至 2017-08-15"],
            ["测试时间范围", "2017-08-16 至 2017-08-31"],
            ["评价指标", "RMSLE，越低越好"],
            ["实验版本", "V1 基础版（V6代码）/ V2 自研增强版（V7代码）/ V3 融合优化版"],
            ["课程截止时间", "2026年6月30日 23:59"],
        ],
        columns=["项目", "内容"],
    )
    add_df_table(doc, meta, "表1 实验基本信息")


def build_full_docx(figures: dict[str, Path]) -> None:
    doc = Document()
    configure_doc(doc)
    add_simple_title(doc)
    add_metadata(doc)

    add_heading(doc, "摘要", 1)
    add_para(doc, "本实验选择 Kaggle 的 Store Sales - Time Series Forecasting 赛题，围绕多门店、多品类、强季节性零售销量预测展开。赛题要求根据 2013-2017 年历史销量、门店信息、商品品类、促销数量、油价、节假日事件和交易量等数据，预测 2017-08-16 至 2017-08-31 共 16 天的销售额。由于数据同时具有时间序列、层级结构、节假日冲击、促销变化和大量低销量/零销量样本，本题不能简单套用普通回归模型，而需要结合业务规律做特征工程和时间顺序验证。")
    add_para(doc, "我的实验过程采用 V1、V2、V3 三阶段迭代。V1 基于 kaggle_notebook_v6.py 完成数据读取、合并、基础特征和树模型 baseline；V2 基于 kaggle_notebook_v7.py 进行自研增强，重点加入 Fourier 周期、趋势项、rank 编码、STL 风格分解、lag/rolling/EWM、促销和节假日交互、多模型融合与后处理；V3 在 V2 的基础上继续做组件化实验、工程稳定性修复、log-space 融合和本地权重搜索，最终 Public RMSLE 达到 0.37927。")
    add_para(doc, "关键词：时间序列预测；特征工程；LightGBM；XGBoost；CatBoost；模型融合；RMSLE。")

    add_heading(doc, "1 题目分析", 1)
    add_heading(doc, "1.1 业务背景与预测对象", 2)
    add_para(doc, "Store Sales 赛题来自厄瓜多尔大型零售商 Favorita。零售销量预测的难点在于需求同时受长期趋势、季节周期、星期效应、节假日、促销、油价和门店结构影响。对商超来说，销量预测直接影响补货计划、库存周转和促销安排，预测偏高会造成积压，预测偏低会造成缺货。")
    add_para(doc, "本题每条样本由 date、store_nbr、family 三个核心字段确定。store_nbr 表示门店，family 表示商品品类，sales 是需要预测的目标值，onpromotion 表示当天该门店该品类的促销商品数量。最终测试集共有 54 个门店、33 个品类、16 个预测日期，因此提交行数为 54 × 33 × 16 = 28,512 行。")
    add_picture(doc, figures["data_relation"], "图1 数据表关系与建模输入", 6.5)

    add_heading(doc, "1.2 赛题难点", 2)
    add_bullet(doc, "多层级结构：销量同时受到门店、城市、州、cluster、商品品类等层级影响，不同门店-品类组合的基准销量差异很大。")
    add_bullet(doc, "强时间依赖：存在周周期、月周期、年周期，以及节前节后、工资日、月末等特殊日期效应。")
    add_bullet(doc, "外部事件冲击：节假日、地震、油价变化、促销数量都会改变短期销量。")
    add_bullet(doc, "零销量和闭店问题：部分组合长期为零，部分门店在特定日期闭店，若不处理会显著影响 RMSLE。")
    add_bullet(doc, "本地验证与榜单不完全一致：Public Leaderboard 只覆盖测试期的一部分真实表现，过度追榜可能造成不稳定。")

    add_heading(doc, "1.3 评价指标 RMSLE", 2)
    add_picture(doc, figures["rmsle"], "图2 RMSLE 指标与建模启发", 6.2)
    add_para(doc, "竞赛使用 RMSLE 作为评价指标。RMSLE 先对真实值和预测值取 log(1+x)，再计算均方根误差。由于 log 变换压缩了极大销量，模型不应只追求大销量样本的绝对误差，而要兼顾不同规模序列上的相对误差。这也是我在 V2/V3 中统一使用 log1p(sales) 作为训练目标、融合时使用 log-space blend 的原因。")
    add_code_block(doc, "RMSLE = sqrt(mean((log(1 + y_pred) - log(1 + y_true)) ** 2))\n训练目标: y_log = log1p(sales)\n提交还原: sales_pred = expm1(pred_log).clip(0, None)")

    add_heading(doc, "2 数据集说明与探索性分析", 1)
    overview = read_csv("dataset_overview.csv")
    add_df_table(doc, overview, "表2 数据集规模概览")
    files_df = pd.DataFrame(
        [
            ["train.csv", "历史训练样本", "id, date, store_nbr, family, sales, onpromotion"],
            ["test.csv", "待预测样本", "id, date, store_nbr, family, onpromotion"],
            ["stores.csv", "门店属性", "city, state, type, cluster"],
            ["oil.csv", "每日油价", "date, dcoilwtico"],
            ["holidays_events.csv", "节假日事件", "type, locale, locale_name, transferred"],
            ["transactions.csv", "门店交易量", "date, store_nbr, transactions"],
        ],
        columns=["文件", "含义", "主要字段"],
    )
    add_df_table(doc, files_df, "表3 数据文件说明")
    add_para(doc, "训练集有 3,000,888 行，测试集有 28,512 行，门店数为 54，品类数为 33。数据时间跨度较长，覆盖多个年度，因此可以提取年周期和长期趋势。但测试集只预测未来 16 天，短期节假日、促销和最近历史 lag 也非常重要。")

    add_picture(doc, figures["monthly"], "图3 月度销售额趋势", 6.4)
    add_para(doc, "月度销售额趋势显示，销售并不是平稳序列，存在长期变化与周期波动。对这类数据，如果只使用静态类别特征，模型很难捕捉未来 16 天的时间位置。因此 V2 中加入了 trend、Fourier、lag、rolling、EWM 等时序特征。")
    add_picture(doc, figures["top_family"], "图4 Top 商品品类销售额", 6.2)
    add_para(doc, "不同商品品类销售规模差异明显，例如 Grocery、Beverages、Produce 等品类贡献较大，而部分品类销量较低。为了让模型学习这种差异，V7 中加入 family_rank、store_rank、store-family 统计特征，以及 family 与 dayofweek/month/quarter 的交互统计。")
    add_picture(doc, figures["store_type"], "图5 门店类型分布", 5.8)
    add_picture(doc, figures["cluster"], "图6 门店 Cluster 分布", 6.3)
    add_para(doc, "门店 type 和 cluster 反映了门店规模与经营模式差异。V7 不仅使用它们作为类别特征，还构造 cluster_rank、cluster × family、cluster × trend 等交互，让模型区分不同门店群体的销售趋势。")
    add_picture(doc, figures["promo_oil"], "图7 销售额、促销数量与油价的月度关系", 6.4)
    add_para(doc, "油价是厄瓜多尔经济环境的重要外部变量，促销数量则直接影响短期销量。EDA 中可以看到这些变量与销售走势并非完全无关，因此在 V2 中分别构造油价移动平均/变化率、促销 lag/rolling 以及 promo × holiday 交互。")
    add_picture(doc, figures["holiday"], "图8 节假日事件类型分布", 5.8)
    add_para(doc, "节假日数据包含 Holiday、Event、Additional、Transfer、Bridge、Work Day 等类型。节假日不仅影响当天销售，也会影响节前备货和节后回落，所以 V7 加入 days_to_holiday 与 days_from_holiday，并保留 national/regional/local 三层节假日信息。")

    add_heading(doc, "3 总体求解思路", 1)
    add_picture(doc, figures["workflow"], "图9 V1/V2/V3 实验迭代流程", 6.5)
    add_para(doc, "整体求解流程可以概括为：先把所有原始表按 date、store_nbr、family 合并成统一样本表；再按照时间顺序构造不泄漏未来的特征；然后使用最后一段时间作为验证集评估模型；最后在全量训练集上重新训练并生成提交文件。实验迭代不是一次性完成，而是从 V1 的可运行 baseline 逐步发展到 V2 的系统特征工程，再进入 V3 的融合和稳定性优化。")
    solution_steps = pd.DataFrame(
        [
            ["1", "数据准备", "读取 train/test/stores/oil/holidays/transactions，统一日期格式和数据类型。"],
            ["2", "特征构造", "生成日历、油价、节假日、促销、lag、rolling、趋势、Fourier、交互统计等特征。"],
            ["3", "时间验证", "按时间切出验证窗口，避免随机切分造成未来信息泄漏。"],
            ["4", "模型训练", "训练 LightGBM、XGBoost、CatBoost，并记录单模型 RMSLE。"],
            ["5", "融合搜索", "在 log-space 搜索多模型/多组件权重，选择验证集更稳的组合。"],
            ["6", "后处理提交", "非负截断、闭店置零、分布检查，生成 submission.csv。"],
        ],
        columns=["步骤", "模块", "具体工作"],
    )
    add_df_table(doc, solution_steps, "表4 具体求解流程")
    add_picture(doc, figures["timeline"], "图10 时间序列验证切分", 6.4)

    add_heading(doc, "4 V1 基础版：从数据到可提交 baseline", 1)
    add_picture(doc, figures["v1_pipeline"], "图16 V1 基础版建模流程", 6.4)
    add_para(doc, "V1 对应 kaggle_notebook_v6.py。这个版本的主要目标是跑通完整流程，而不是一次性追求最高分。V1 完成了数据读取、字段清洗、缺失油价填充、节假日透视表、门店信息合并和基础滞后特征构造，并尝试使用 XGBoost、LightGBM 等树模型建模。")
    add_picture(doc, figures["v1_features"], "图17 V1 基础特征范围", 6.1)
    add_para(doc, "V1 的处理逻辑中，先将 National、Regional、Local holiday 分别整理成特征表，再与主表按照 date、city/state 等键合并；油价使用前向/后向填充；门店属性作为类别变量进入模型；最后构造一部分 lag 特征用于未来 16 天预测。这个版本让我明确了题目的基本数据结构，也验证了树模型适合处理该类结构化时序特征。")
    add_picture(doc, figures["v1_diag"], "图18 V1 结果诊断与 V2 改进方向", 6.2)
    add_para(doc, "V1 的不足也比较明显：lag 特征较少，周期和趋势表达不足；节假日和促销的交互没有充分利用；验证方式和测试期分布不够贴近；模型融合和后处理较弱。因此，V2 的重点转向系统性特征工程和更稳健的验证流程。")

    add_heading(doc, "5 V2 自研增强版：V7 建模过程", 1)
    add_para(doc, "V2 是本实验的主要自主建模阶段，对应 kaggle_notebook_v7.py。V7 的设计思路是：不把销量预测当作普通表格回归，而是把每个 store-family 看作一条时间序列，再把时间序列规律、门店层级、商品品类、油价、节假日和促销信息统一编码成树模型可学习的特征。")
    add_picture(doc, figures["v2_upgrade"], "图19 V2 相比 V1 的核心升级", 6.6)
    add_picture(doc, figures["v7_features"], "图11 V7 自研特征工程结构", 6.6)

    add_heading(doc, "5.1 数据合并与内存优化", 2)
    add_para(doc, "V7 首先将 train 与 test 合并，增加 is_train 标记，并将 test 的 sales 置为缺失。这样做的好处是特征工程可以在完整日期范围上统一完成，避免训练集和测试集特征列不一致。合并 stores 后，数据粒度保持为 date × store_nbr × family。")
    add_para(doc, "由于训练集超过 300 万行，V7 增加 reduce_mem_usage，将 float64 转为 float32，将 int64 转为 int32 或更小整数类型，以降低内存占用。这个工程细节很重要，因为后续会生成大量 lag、rolling 和交互特征，若不压缩类型，很容易在 Kaggle 或 AutoDL 环境中内存不足。")

    add_heading(doc, "5.2 目标变换与异常值处理", 2)
    add_para(doc, "评价指标 RMSLE 与 log 误差一致，因此 V7 使用 sales_log = log1p(sales) 作为主要训练目标。对销量中的极端大值，V7 按 store-family 组合进行 99.5% 分位数裁剪，并保留 is_clipped 标记。这样既能减弱异常峰值对模型训练的干扰，又不会完全删除这些样本的信息。")
    add_code_block(doc, "sales_log = np.log1p(sales.clip(lower=0))\nclip_upper = sales.groupby([store_nbr, family]).transform(lambda x: x.quantile(0.995))\nis_clipped = sales > clip_upper\nsales = np.minimum(sales, clip_upper).astype('float32')")

    add_heading(doc, "5.3 日历、油价和节假日特征", 2)
    add_para(doc, "日期特征包括 year、month、day、dayofweek、weekofyear、quarter、is_weekend、is_month_end、is_payday 等。工资日特征针对 15 日和月末设计，因为零售消费可能在发薪日前后出现变化。")
    add_para(doc, "油价特征先在完整日期序列上插值，再构造 oil_ma7、oil_ma30、oil_ma90、oil_chg7、oil_chg30。这样模型不仅看到当天油价，也能看到短期和中期油价趋势。")
    add_para(doc, "节假日特征分为 National、Regional、Local 三层，并处理 Work Day 和 transferred holiday。V7 构造 is_holiday、is_reg_holiday、is_local_holiday、holiday_type_enc、days_to_holiday、days_from_holiday 等特征，用于捕捉节前提前购买、节后回落等现象。")

    add_heading(doc, "5.4 Fourier、趋势与周期特征", 2)
    add_para(doc, "V7 引入 Fourier 周期特征，用 sin/cos 表达 3.5、7、30、365 等周期。周周期可捕捉工作日/周末差异，月周期对应月初月末消费模式，年周期对应季节性和年度节日。相比只加入 dayofweek 或 month，Fourier 特征能更连续地表达周期位置。")
    add_para(doc, "趋势特征包括 trend_linear、trend_squared，并进一步构造 family_rank × trend、cluster_rank × trend。这是因为不同品类和不同门店群体的长期走势并不相同：有些品类可能增长，有些品类可能衰退，直接用全局 trend 会掩盖这种差异。")

    add_heading(doc, "5.5 lag、rolling、EWM 与 momentum", 2)
    add_picture(doc, figures["v2_lag"], "图20 V2 lag 设计与防止未来信息泄漏", 6.5)
    add_para(doc, "时间序列预测最重要的信息来自历史销量。V7 以 store_nbr-family 为分组，构造 lag_16、lag_17、lag_18、lag_19、lag_20、lag_21、lag_28、lag_35、lag_56、lag_84、lag_112、lag_182、lag_364。选择从 16 开始，是为了避免预测未来 16 天时使用不可获得的未来真实销量。")
    add_para(doc, "在 lag 基础上，V7 构造 rolling mean、rolling std、EWM 和 momentum。rolling mean 描述近期平均水平，rolling std 描述波动程度，EWM 强调越近的历史越重要，momentum 描述销量相对近期均值的变化方向。促销变量 onpromotion 也构造对应 lag 和 rolling 特征。")
    add_code_block(doc, "for lag in [16,17,18,19,20,21,28,35,56,84,112,182,364]:\n    data[f'lag_{lag}'] = data.groupby(['store_nbr','family'])['sales_log'].shift(lag)\n\nrolling_mean_28 = group.sales_log.shift(16).rolling(28).mean()\nrolling_std_28  = group.sales_log.shift(16).rolling(28).std()\newm_28 = group.sales_log.shift(16).ewm(span=28).mean()")

    add_heading(doc, "5.6 交互统计、rank 编码与 STL 风格分解", 2)
    add_para(doc, "V7 还构造多组交互统计，例如 store-family 均值/中位数/std，city-family、state-family、cluster-family、store-month、family-dayofweek、family-month、family-quarter 等。这些统计特征相当于给模型提供不同层级的历史先验。")
    add_para(doc, "rank 编码包括 family_rank、store_rank、cluster_rank，表示不同品类、门店和 cluster 的相对销售能力。相比直接使用类别编码，rank 特征提供了有序信息。STL 风格分解则通过 lag_364 和滚动均值近似 trend、seasonal、residual、seasonal_strength，使模型更容易区分长期水平、季节偏移和异常波动。")

    add_heading(doc, "5.7 模型训练与验证", 2)
    add_picture(doc, figures["ensemble"], "图12 模型训练、融合与后处理流程", 6.5)
    add_para(doc, "V7 使用按时间顺序的 31 天验证窗口，而不是随机切分。随机切分会让模型在训练中看到未来日期附近的样本，导致验证分数虚高；时间窗口验证更接近真实提交场景。")
    add_para(doc, "模型方面，V7 使用 LightGBM、XGBoost 和 CatBoost。LightGBM 通过 Optuna 搜索参数，并训练多个 seed 取平均；XGBoost 使用多组 max_depth、learning_rate、subsample、colsample_bytree 配置；CatBoost 作为对类别特征较友好的补充模型。所有模型都预测 log-space 的 sales_log。")
    add_para(doc, "融合时先在验证集上记录各模型 RMSLE，再搜索权重组合。设不同模型的 log 预测为 p_i，权重为 w_i，最终 log 预测为 sum(w_i * p_i)，再通过 expm1 还原销量。这样与 RMSLE 指标保持一致，也能减少极端预测对融合结果的影响。")

    add_heading(doc, "5.8 后处理策略", 2)
    add_picture(doc, figures["v2_post"], "图21 V2 后处理策略", 6.2)
    add_para(doc, "后处理主要包括三类：第一，所有预测值 clip 到非负；第二，对闭店或长期零销量组合进行置零；第三，构造季节性预测并与模型预测按 family 自适应融合。V7 还尝试过 pseudo-labeling，但考虑到伪标签容易带来不稳定，后续稳定版本中把它作为对照实验而不是默认策略。")
    add_para(doc, "V2 最终 Public Score 约为 0.40。虽然距离最终最好成绩还有差距，但它完成了从 baseline 到系统建模的关键跨越，也是后续 V3 融合优化的技术基础。")

    add_heading(doc, "6 V3 融合优化版：从 V2 到最终提交", 1)
    add_para(doc, "V2 之后继续堆叠复杂特征和模型，训练成本会明显上升，收益也开始变小。因此 V3 的目标不是推翻 V2，而是在 V2 的经验上做更稳的工程化和融合优化：把不同训练窗口、不同算法族、不同提交候选拆成组件，统一成可比较、可调权重、可分布检查的实验对象。")
    add_picture(doc, figures["v3_arch"], "图22 V3 组件化融合结构", 6.4)
    add_para(doc, "V3 的具体工作包括：整理 AutoDL 可运行脚本，统一输入输出路径；修复库版本兼容问题；生成多组候选提交；比较 linear blend 与 log blend；对四个主要组件进行权重搜索；检查最终提交的均值、中位数、零值数量和最大预测值，避免融合后出现异常分布。")
    add_picture(doc, figures["v3_weight_process"], "图23 V3 权重搜索过程示意", 6.2)
    add_df_table(doc, read_csv("best_component_weights.csv"), "表5 V3 最终四组件权重")
    add_picture(doc, figures["weights"], "图13 V3 四组件权重")
    add_para(doc, "最终权重为 lgb_full 32.5%、lgb_2015 32.5%、xgb_full 35.0%、xgb_2015 0%。虽然第四个组件在本地搜索中权重为 0，但它仍然有实验价值，因为它证明了并不是组件越多越好，融合应该以验证效果和分布稳定性为准。")
    add_code_block(doc, "pred_final = expm1(\n    0.325 * log1p(pred_lgb_full)\n  + 0.325 * log1p(pred_lgb_2015)\n  + 0.350 * log1p(pred_xgb_full)\n  + 0.000 * log1p(pred_xgb_2015)\n)")
    add_picture(doc, figures["submission_dist"], "图14 候选提交预测分布检查", 6.2)

    add_heading(doc, "7 实验结果与排名分析", 1)
    raw_scores = read_csv("score_progression.csv")
    scores = pd.DataFrame(
        {
            "阶段": ["V2 自研增强模型", "V3-1 时序组件", "V3-2 分层融合", "V3-3 树模型组件", "V3-4 log融合", "V3-5 轻量LGBM组件", "V3-6 本地权重搜索"],
            "Public RMSLE": raw_scores["public_score"].astype(float).round(5),
            "说明": ["V7主模型", "候选组件", "候选组件", "候选组件", "融合提升", "小幅提升", "最终最好"],
        }
    )
    add_df_table(doc, scores, "表6 分数迭代记录")
    add_picture(doc, figures["score"], "图15 Public Score 迭代过程", 6.5)
    add_picture(doc, figures["v3_waterfall"], "图24 V3 成绩推进瀑布图", 6.5)
    add_para(doc, "从结果看，V2 自研增强模型约为 0.40，说明系统特征工程已经明显优于基础 baseline。V3 通过组件化实验和 log-space 融合，将成绩推进到 0.37927。虽然后期提升看起来只有 0.0001 量级，但在排行榜前段，这种微小差异也会改变名次。")
    add_para(doc, "Kaggle 页面截图显示，在 Public Score 0.37932 时账号 Simoni Fretta 排名约第 298 名，后续最终成绩提升到 0.37927。最终名次以 Kaggle 页面显示为准。")

    add_heading(doc, "8 结果讨论", 1)
    add_heading(doc, "8.1 为什么 V7 分数没有直接到 0.38 以下", 2)
    add_para(doc, "V7 的特征工程较完整，但它仍然是一个单套主流程。Store Sales 这类赛题的榜单前段通常非常依赖细粒度验证、特定窗口建模和提交融合。V7 中的特征数量多、模型重，局部改动很容易造成验证集和榜单不一致，例如节假日 merge、异常值 clip、pseudo-labeling 等都可能改变分布。")
    add_heading(doc, "8.2 为什么 V3 选择 log blend", 2)
    add_para(doc, "RMSLE 本质上衡量 log 空间误差，因此在原始 sales 上做线性平均，不一定符合指标。log blend 先把不同组件预测转到 log1p 空间再加权，可以减少大销量样本对融合的支配，也能让低销量样本的相对误差更稳定。实际实验中，log blend 相比 linear blend 更适合作为最终融合方式。")
    add_heading(doc, "8.3 本地验证与 Public 分数的关系", 2)
    add_para(doc, "本地验证用于筛选方向，但不能完全等同于 Public Leaderboard。时间序列任务中，不同年份、不同节假日位置、不同促销强度都会让验证窗口和测试窗口分布不同。因此我在 V3 中没有只看单一验证分数，还检查了提交文件的均值、中位数、零销量行数和最大值，避免出现明显分布漂移。")

    add_heading(doc, "9 代码与输出文件说明", 1)
    code_df = pd.DataFrame(
        [
            ["kaggle_notebook_v6.py", "V1 基础版", "完成基础数据处理、合并、初步 lag 和树模型 baseline。"],
            ["kaggle_notebook_v7.py", "V2 自研增强版", "完整特征工程、多模型集成、验证集融合和后处理。"],
            ["store_sales_xiewenwei_clean.py / 融合脚本", "V3 融合优化版", "组件化实验、log blend、权重搜索和最终提交。"],
            ["report_assets/", "报告素材", "保存图表、表格和本 Word 报告。"],
        ],
        columns=["文件/目录", "所属阶段", "作用"],
    )
    add_df_table(doc, code_df, "表7 代码与输出文件")
    add_para(doc, "主要提交文件包括 submission.csv、submission_no_pseudo.csv、submission_lgbm_final.csv、submission_xgb_final.csv 以及最终融合提交文件。报告中的图表素材来自训练数据统计、提交文件统计和实验记录。")

    add_heading(doc, "10 总结与改进方向", 1)
    add_para(doc, "本次实验完整覆盖了 Kaggle 竞赛流程：注册参赛、下载数据、理解题目、编写算法、提交结果、查看排名、继续改进。V1 让我完成从数据到提交的闭环；V2 是主要自主建模成果，围绕 Store Sales 的业务结构做了较系统的特征工程；V3 在此基础上通过融合和工程稳定性优化，把成绩提升到 0.37927。")
    add_para(doc, "后续仍可改进的方向包括：构造多窗口时间序列交叉验证，减少对单一验证窗口和 Public Leaderboard 的依赖；对闭店、低销量和新品类样本单独建模；进一步分析节假日和促销的交互窗口；把 V3 中验证有效的融合策略沉淀回 V7 主代码，使自研版本更加稳定和简洁。")
    add_para(doc, "总体来看，本实验不仅得到较好的榜单成绩，也让我熟悉了结构化时间序列预测的完整方法：从 EDA 到特征工程，从模型训练到融合，从本地验证到提交分布检查。")

    doc.save(OUT_FULL_DOCX)


def main() -> None:
    WORD_FIG_DIR.mkdir(parents=True, exist_ok=True)
    figures = {
        "monthly": line_chart_monthly(),
        "top_family": bar_chart_top_family(),
        "store_type": bar_chart_store_type(),
        "workflow": workflow_chart(),
        "weights": component_weights_chart(),
        "submission_dist": submission_distribution_chart(),
        "score": score_progression_chart(),
        "data_relation": data_relationship_chart(),
        "timeline": validation_timeline_chart(),
        "rmsle": rmsle_explanation_chart(),
        "v7_features": v7_feature_chart(),
        "ensemble": ensemble_pipeline_chart(),
        "promo_oil": promo_oil_chart(),
        "holiday": holiday_bar_chart(),
        "cluster": cluster_bar_chart(),
        "v1_pipeline": v1_baseline_pipeline_chart(),
        "v1_features": v1_feature_scope_chart(),
        "v1_diag": v1_problem_diagnosis_chart(),
        "v2_upgrade": v2_upgrade_map_chart(),
        "v2_lag": v2_lag_leakage_chart(),
        "v2_post": v2_postprocess_chart(),
        "v3_arch": v3_component_architecture_chart(),
        "v3_weight_process": v3_weight_search_process_chart(),
        "v3_waterfall": v3_score_waterfall_chart(),
    }
    build_full_docx(figures)
    print(f"Saved full report: {OUT_FULL_DOCX}")
    print(f"Figures: {WORD_FIG_DIR}")


if __name__ == "__main__":
    main()
