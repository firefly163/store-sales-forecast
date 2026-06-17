from __future__ import annotations

from pathlib import Path
import math
import textwrap

import pandas as pd
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSET_DIR = ROOT / "report_assets"
TABLE_DIR = ASSET_DIR / "tables"
WORD_FIG_DIR = ASSET_DIR / "word_figures"
OUT_DOCX = ASSET_DIR / "Store_Sales_实验报告.docx"
OUT_MD = ASSET_DIR / "Store_Sales_实验报告.md"

BLUE = (31, 78, 121)
LIGHT_BLUE = (221, 235, 247)
GRAY = (90, 90, 90)
LIGHT_GRAY = (245, 247, 250)
ORANGE = (237, 125, 49)
GREEN = (112, 173, 71)


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLE_DIR / name)


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(exist_ok=True)
    WORD_FIG_DIR.mkdir(exist_ok=True)


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        try:
            if Path(path).exists():
                return ImageFont.truetype(path, size)
        except Exception:
            pass
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def wrapped_lines(text: str, width: int) -> list[str]:
    lines: list[str] = []
    for part in str(text).split("\n"):
        lines.extend(textwrap.wrap(part, width=width) or [""])
    return lines


def draw_title(draw: ImageDraw.ImageDraw, title: str, w: int) -> None:
    font = get_font(32, True)
    tw, _ = text_size(draw, title, font)
    draw.text(((w - tw) / 2, 28), title, fill=BLUE, font=font)


def nice_max(v: float) -> float:
    if v <= 0:
        return 1
    exp = math.floor(math.log10(v))
    base = 10 ** exp
    for m in [1, 2, 5, 10]:
        if v <= m * base:
            return m * base
    return 10 * base


def line_chart_monthly() -> Path:
    df = read_csv("monthly_sales.csv")
    x_col = df.columns[0]
    y_col = "sales" if "sales" in df.columns else df.select_dtypes("number").columns[0]
    df = df.copy()
    df[x_col] = df[x_col].astype(str)
    vals = df[y_col].astype(float).values

    out = WORD_FIG_DIR / "monthly_sales_trend.png"
    w, h = 1200, 650
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "月度销售额趋势", w)

    left, right, top, bottom = 90, 40, 95, 95
    plot_w, plot_h = w - left - right, h - top - bottom
    ymin, ymax = 0, nice_max(float(vals.max()))

    draw.rectangle((left, top, left + plot_w, top + plot_h), outline=(210, 210, 210), width=2)
    axis_font = get_font(18)
    for i in range(5):
        y = top + plot_h - i * plot_h / 4
        val = ymin + i * (ymax - ymin) / 4
        draw.line((left, y, left + plot_w, y), fill=(230, 230, 230), width=1)
        draw.text((8, y - 10), f"{val/1e6:.1f}M", fill=GRAY, font=axis_font)

    pts = []
    n = len(vals)
    for i, v in enumerate(vals):
        x = left + i * plot_w / max(1, n - 1)
        y = top + plot_h - (float(v) - ymin) / (ymax - ymin) * plot_h
        pts.append((x, y))
    if len(pts) > 1:
        draw.line(pts, fill=BLUE, width=4)
    for x, y in pts[:: max(1, len(pts) // 20)]:
        draw.ellipse((x - 3, y - 3, x + 3, y + 3), fill=ORANGE)

    for i in range(0, n, max(1, n // 8)):
        x = left + i * plot_w / max(1, n - 1)
        label = df[x_col].iloc[i][:7]
        tw, _ = text_size(draw, label, axis_font)
        draw.text((x - tw / 2, top + plot_h + 18), label, fill=GRAY, font=axis_font)

    draw.text((left, h - 35), "说明：销售额存在明显趋势、季节性和节假日/促销扰动，适合使用时序特征与集成模型。", fill=GRAY, font=get_font(18))
    img.save(out, quality=95)
    return out


def bar_chart_top_family() -> Path:
    df = read_csv("family_sales_total.csv")
    family_col = "family" if "family" in df.columns else df.columns[0]
    value_col = "sales" if "sales" in df.columns else df.select_dtypes("number").columns[0]
    df = df.sort_values(value_col, ascending=False).head(12)

    out = WORD_FIG_DIR / "top_family_sales.png"
    w, h = 1200, 700
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "Top 商品品类销售额", w)

    left, top, bar_h, gap = 280, 110, 30, 14
    max_v = float(df[value_col].max())
    font = get_font(20)
    small = get_font(17)
    for i, row in enumerate(df.itertuples(index=False)):
        family = str(getattr(row, family_col))
        val = float(getattr(row, value_col))
        y = top + i * (bar_h + gap)
        bw = int((w - left - 130) * val / max_v)
        draw.text((30, y + 3), family[:24], fill=GRAY, font=font)
        draw.rounded_rectangle((left, y, left + bw, y + bar_h), radius=8, fill=BLUE if i < 5 else (91, 155, 213))
        draw.text((left + bw + 12, y + 4), f"{val/1e6:.1f}M", fill=GRAY, font=small)

    img.save(out, quality=95)
    return out


def bar_chart_store_type() -> Path:
    df = read_csv("store_type_counts.csv")
    label_col = df.columns[0]
    value_col = df.select_dtypes("number").columns[0]

    out = WORD_FIG_DIR / "store_type_counts.png"
    w, h = 1000, 560
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "门店类型数量分布", w)

    left, base_y = 120, 460
    plot_w, plot_h = 760, 320
    max_v = nice_max(float(df[value_col].max()))
    bar_w = plot_w / len(df) * 0.62
    font = get_font(22)
    small = get_font(18)

    draw.line((left, base_y, left + plot_w, base_y), fill=(180, 180, 180), width=2)
    for i, row in enumerate(df.itertuples(index=False)):
        label = str(getattr(row, label_col))
        val = float(getattr(row, value_col))
        x = left + i * plot_w / len(df) + plot_w / len(df) * 0.19
        bh = plot_h * val / max_v
        draw.rounded_rectangle((x, base_y - bh, x + bar_w, base_y), radius=8, fill=ORANGE if i % 2 else BLUE)
        tw, _ = text_size(draw, label, font)
        draw.text((x + bar_w / 2 - tw / 2, base_y + 18), label, fill=GRAY, font=font)
        val_txt = str(int(val))
        tw, _ = text_size(draw, val_txt, small)
        draw.text((x + bar_w / 2 - tw / 2, base_y - bh - 28), val_txt, fill=GRAY, font=small)

    img.save(out, quality=95)
    return out


def score_progression_chart() -> Path:
    df = read_csv("score_progression.csv")
    df = df.copy()
    df["public_score"] = df["public_score"].astype(float)

    out = WORD_FIG_DIR / "score_progression.png"
    w, h = 1250, 720
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "提交成绩迭代过程（Public RMSLE，越低越好）", w)

    left, right, top, bottom = 95, 55, 105, 155
    plot_w, plot_h = w - left - right, h - top - bottom
    vals = df["public_score"].values
    ymin = math.floor((float(vals.min()) - 0.005) * 1000) / 1000
    ymax = math.ceil((float(vals.max()) + 0.005) * 1000) / 1000

    draw.rectangle((left, top, left + plot_w, top + plot_h), outline=(210, 210, 210), width=2)
    axis_font = get_font(17)
    for i in range(6):
        y = top + i * plot_h / 5
        val = ymax - i * (ymax - ymin) / 5
        draw.line((left, y, left + plot_w, y), fill=(232, 232, 232), width=1)
        draw.text((22, y - 10), f"{val:.3f}", fill=GRAY, font=axis_font)

    pts = []
    n = len(df)
    for i, v in enumerate(vals):
        x = left + i * plot_w / max(1, n - 1)
        y = top + (ymax - float(v)) / (ymax - ymin) * plot_h
        pts.append((x, y))
    draw.line(pts, fill=BLUE, width=4)
    for i, (x, y) in enumerate(pts):
        color = GREEN if i == n - 1 else ORANGE
        draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill=color)
        draw.text((x - 35, y - 36), f"{vals[i]:.5f}", fill=GRAY, font=axis_font)

    labels = [
        "V2 自研",
        "V3-1\n时序组件",
        "V3-2\n分层融合",
        "V3-3\n树模型组件",
        "V3-4\nlog融合",
        "V3-5\n轻量LGBM",
        "V3-6\n本地权重",
    ]
    for i, label in enumerate(labels[:n]):
        x = left + i * plot_w / max(1, n - 1)
        y = top + plot_h + 22
        for j, line in enumerate(wrapped_lines(label, 8)):
            tw, _ = text_size(draw, line, axis_font)
            draw.text((x - tw / 2, y + j * 22), line, fill=GRAY, font=axis_font)

    img.save(out, quality=95)
    return out


def component_weights_chart() -> Path:
    df = read_csv("best_component_weights.csv")
    out = WORD_FIG_DIR / "component_weights.png"
    w, h = 1000, 560
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "最终本地权重搜索结果", w)

    left, base_y = 120, 440
    plot_w, plot_h = 760, 300
    font = get_font(20)
    small = get_font(18)
    colors = [BLUE, (91, 155, 213), ORANGE, (180, 180, 180)]
    for i, row in df.iterrows():
        label = str(row["component"])
        val = float(row["weight"])
        x = left + i * plot_w / len(df) + 35
        bw = plot_w / len(df) - 70
        bh = plot_h * val
        draw.rounded_rectangle((x, base_y - bh, x + bw, base_y), radius=8, fill=colors[i % len(colors)])
        txt = f"{val:.1%}"
        tw, _ = text_size(draw, txt, small)
        draw.text((x + bw / 2 - tw / 2, base_y - bh - 28), txt, fill=GRAY, font=small)
        for j, line in enumerate(label.split("_")):
            tw, _ = text_size(draw, line, font)
            draw.text((x + bw / 2 - tw / 2, base_y + 16 + j * 23), line, fill=GRAY, font=font)
    draw.text((left, h - 42), "融合采用 log-space：先对预测值取 log1p 后加权，再通过 expm1 还原。", fill=GRAY, font=small)

    img.save(out, quality=95)
    return out


def workflow_chart() -> Path:
    out = WORD_FIG_DIR / "experiment_workflow.png"
    w, h = 1300, 620
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "实验迭代流程", w)

    boxes = [
        ("V1 基础版\nEDA + 数据合并\n油价/假日/门店特征\nXGB/LGB 基线", 50, 145, BLUE),
        ("V2 自研增强版\nFourier/趋势/rank\nlag/rolling/STL特征\nLGB+XGB+CatBoost", 355, 145, ORANGE),
        ("V3 融合优化版\n组件化实验\nlog blend\n本地权重搜索", 690, 145, GREEN),
        ("最终提交\nPublic RMSLE 0.37927\n生成报告素材\n总结排名与反思", 1015, 145, (112, 48, 160)),
    ]
    title_font = get_font(25, True)
    body_font = get_font(20)
    for text, x, y, color in boxes:
        draw.rounded_rectangle((x, y, x + 245, y + 260), radius=20, fill=LIGHT_GRAY, outline=color, width=4)
        parts = text.split("\n")
        tw, _ = text_size(draw, parts[0], title_font)
        draw.text((x + 122 - tw / 2, y + 25), parts[0], fill=color, font=title_font)
        for j, line in enumerate(parts[1:]):
            tw, _ = text_size(draw, line, body_font)
            draw.text((x + 122 - tw / 2, y + 84 + j * 38), line, fill=GRAY, font=body_font)
    for x in [302, 637, 962]:
        draw.line((x, 275, x + 42, 275), fill=(110, 110, 110), width=4)
        draw.polygon([(x + 42, 275), (x + 26, 265), (x + 26, 285)], fill=(110, 110, 110))

    draw.text((60, 480), "迭代思路：从可运行基线到自研特征增强，再到融合策略、工程稳定性和提交优化。", fill=GRAY, font=get_font(21))
    img.save(out, quality=95)
    return out


def submission_distribution_chart() -> Path:
    df = read_csv("submission_distribution_stats.csv")
    out = WORD_FIG_DIR / "submission_distribution.png"
    w, h = 1200, 580
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "提交文件预测分布对比", w)

    left, top = 80, 125
    col_w = 340
    font = get_font(22, True)
    small = get_font(19)
    for i, row in df.iterrows():
        x = left + i * col_w
        draw.rounded_rectangle((x, top, x + 300, top + 330), radius=18, fill=LIGHT_GRAY, outline=(205, 205, 205), width=2)
        title = str(row["submission"])
        tw, _ = text_size(draw, title, font)
        draw.text((x + 150 - tw / 2, top + 24), title, fill=BLUE if i < 2 else GREEN, font=font)
        lines = [
            f"rows: {int(row['rows']):,}",
            f"mean sales: {float(row['mean_sales']):.2f}",
            f"median: {float(row['median_sales']):.2f}",
            f"zero rows: {int(row['zero_rows']):,}",
            f"max sales: {float(row['max_sales']):.1f}",
        ]
        for j, line in enumerate(lines):
            draw.text((x + 40, top + 92 + j * 43), line, fill=GRAY, font=small)
    draw.text((left, h - 50), "分布检查用于确认融合后没有出现异常负值、异常全零或极端预测漂移。", fill=GRAY, font=small)
    img.save(out, quality=95)
    return out


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    set_run_font(run, 9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: tuple[int, int, int] | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 16 if level == 1 else 13, True, BLUE if level == 1 else (45, 45, 45))


def add_para(doc: Document, text: str, size: float = 10.5, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size, bold)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    set_run_font(run, 10.5)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 9, False, GRAY)


def add_picture(doc: Document, path: Path, caption: str, width_inches: float = 6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    add_caption(doc, caption)


def add_df_table(doc: Document, df: pd.DataFrame, title: str | None = None, max_rows: int | None = None) -> None:
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, 10.5, True, BLUE)
    show = df.copy()
    if max_rows is not None:
        show = show.head(max_rows)
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(show.columns):
        set_cell_text(hdr[i], col, True)
        set_cell_shading(hdr[i], "D9EAF7")
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(show.columns):
            val = row[col]
            if isinstance(val, float):
                text = f"{val:.5f}" if abs(val) < 1 else f"{val:,.2f}"
            else:
                text = str(val)
            set_cell_text(cells[i], text)
    doc.add_paragraph()


def add_metadata_table(doc: Document) -> None:
    data = [
        ("竞赛题目", "Store Sales - Time Series Forecasting"),
        ("竞赛平台", "Kaggle"),
        ("课程实验时间", "2026年6月中旬至6月下旬"),
        ("课程提交截止", "2026年6月30日 23:59"),
        ("预测任务", "预测 2017-08-16 至 2017-08-31 各门店-品类销售额"),
        ("评价指标", "RMSLE，数值越低代表排名越好"),
        ("最终 Public Score", "0.37927"),
        ("排名情况", "0.37932 时页面截图排名约第 298 名；最终以 Kaggle 页面为准"),
    ]
    table = doc.add_table(rows=len(data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (k, v) in enumerate(data):
        set_cell_text(table.rows[i].cells[0], k, True)
        set_cell_shading(table.rows[i].cells[0], "D9EAF7")
        set_cell_text(table.rows[i].cells[1], v)
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)


def build_markdown() -> None:
    text = """# Store Sales 时间序列预测实验报告

本文档的 Word 版本已自动生成。报告采用 V1/V2/V3 的迭代结构：V1 对应 V6 基础版，V2 对应 V7 自研特征增强版，V3 对应最终融合优化版。最终 Public RMSLE 为 0.37927。
"""
    OUT_MD.write_text(text, encoding="utf-8")


def build_docx(figures: dict[str, Path]) -> None:
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Store Sales 时间序列预测实验报告")
    set_run_font(run, 22, True, BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Kaggle Store Sales - Time Series Forecasting")
    set_run_font(run, 12, False, GRAY)

    add_metadata_table(doc)

    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "本实验选择 Kaggle 的 Store Sales 时间序列预测赛题，目标是在给定历史销售、门店、油价、节假日和促销信息的基础上，预测未来 16 天每个门店-商品品类组合的销量。实验过程采用逐步迭代方式：V1 完成基础数据处理和可运行基线；V2 在此基础上重点增强时序、节假日、趋势、促销和交互特征，并使用多模型集成；V3 则延续 V2 的建模思路，重点优化组件拆分、log-space 融合、本地权重搜索和提交稳定性，最终 Public RMSLE 达到 0.37927。",
    )
    add_para(
        doc,
        "本报告重点说明自研 V7 的建模思路和特征工程，并将后续融合优化作为成绩提升过程进行总结。整体实验体现了从探索性分析、特征构建、模型训练、提交评估到迭代改进的完整流程。",
    )

    add_heading(doc, "1 赛题任务说明", 1)
    add_para(
        doc,
        "Store Sales - Time Series Forecasting 要求参赛者根据 Corporation Favorita 的历史销售数据，预测测试集 2017-08-16 至 2017-08-31 期间各 store_nbr 与 family 组合的销售额。每条记录对应一个具体日期、门店和商品品类，提交文件需要给出 id 与 sales 两列。",
    )
    add_para(
        doc,
        "竞赛评价指标为 RMSLE。该指标对预测值和真实值取 log1p 后计算均方根误差，能够降低极大销量样本对总体误差的支配，同时对低销量样本和零销量样本更敏感。因此，本题不能只追求大店大品类的拟合，也要处理闭店、低销量、节假日扰动等细节。",
    )
    add_para(doc, "RMSLE = sqrt( mean( (log(1 + y_pred) - log(1 + y_true))^2 ) )")

    add_heading(doc, "2 数据集分析", 1)
    overview = read_csv("dataset_overview.csv")
    add_df_table(doc, overview, "表1 数据集基本规模")
    add_para(
        doc,
        "训练集包含 3,000,888 条记录，测试集包含 28,512 条记录；门店数量为 54，商品品类为 33，因此核心预测对象为 1,782 条门店-品类时间序列。训练时间跨度从 2013-01-01 到 2017-08-15，测试集为紧接其后的 16 天。",
    )
    add_picture(doc, figures["monthly"], "图1 月度销售趋势")
    add_para(
        doc,
        "从月度趋势可以看出，销售额存在明显的长期变化和周期性波动。单纯使用日期字段难以刻画这种规律，因此后续版本中引入了 lag、rolling、EWM、Fourier、趋势项和节假日距离等特征。",
    )
    add_picture(doc, figures["top_family"], "图2 Top 商品品类销售额")
    add_picture(doc, figures["store_type"], "图3 门店类型数量分布")
    add_para(
        doc,
        "商品品类和门店结构差异明显，不同 family、store、cluster 的销售水平并不相同。这个结论直接影响后续特征工程：模型需要同时学习门店维度、品类维度以及二者交互后的销售规律。",
    )

    add_heading(doc, "3 实验迭代路线", 1)
    add_picture(doc, figures["workflow"], "图4 V1/V2/V3 实验迭代流程", 6.6)
    iter_df = pd.DataFrame(
        [
            ["V1", "kaggle_notebook_v6.py", "基础版", "完成 EDA、数据合并、油价/节假日/门店信息处理，使用传统机器学习模型形成可提交 baseline。", "阶段性基线"],
            ["V2", "kaggle_notebook_v7.py", "自研增强版", "围绕时序预测做系统特征工程，加入 Fourier、趋势、rank、STL 风格分解、lag/rolling、交互统计和多模型集成。", "约 0.40"],
            ["V3", "最终融合脚本", "融合优化版", "在 V2 特征与集成经验基础上，进一步做组件化实验、log-space 融合、本地权重搜索和提交文件分布检查。", "0.37927"],
        ],
        columns=["版本", "对应代码", "定位", "主要改进", "Public Score"],
    )
    add_df_table(doc, iter_df, "表2 实验版本迭代说明")

    add_heading(doc, "4 V1 基础版方法", 1)
    add_para(
        doc,
        "V1 对应 kaggle_notebook_v6.py，主要目标不是一次性追求最高分，而是先把完整建模链路跑通。该版本完成了 train/test、stores、oil、holidays_events、transactions 等表的读取与合并，并对字符串字段、缺失油价、假日类型等进行基础处理。",
    )
    add_para(
        doc,
        "在特征方面，V1 已经开始使用日期、门店、商品品类、油价和节假日信息，并尝试构造滞后特征。模型方面主要使用 XGBoost、LightGBM 以及部分 sklearn 模型进行对比，验证了树模型在该类结构化时间序列任务上的有效性。V1 的意义在于建立了可提交的基线，也暴露出两个问题：一是特征对复杂季节性刻画不足，二是验证集和测试集的时间分布需要更严格对齐。",
    )

    add_heading(doc, "5 V2 自研增强版：V7", 1)
    add_para(
        doc,
        "V2 是本次实验中最核心的自研版本，对应 kaggle_notebook_v7.py。该版本保留了 V1 的完整训练与提交流程，但将重点放在 Store Sales 赛题本身的业务规律上：门店-品类销量具有强烈的周周期、月周期、年周期、促销波动、节假日冲击和闭店零销量现象。V7 因此从特征、验证、模型和后处理四个层面进行了增强。",
    )

    add_heading(doc, "5.1 数据预处理与异常值", 2)
    add_para(
        doc,
        "V7 首先对数据类型进行内存优化，将 float64、int64 尽量转换为 float32、int32，以便在 Kaggle 或 AutoDL 环境中稳定训练。针对销量中的极端值，V7 按 store-family 组合进行 99.5% 分位数裁剪，减少异常大单对 log-space 模型训练的影响，同时保留 is_clipped 标记让模型知道该样本曾经被处理。",
    )

    add_heading(doc, "5.2 日历、节假日与油价特征", 2)
    add_para(
        doc,
        "赛题包含油价和节假日外部信息。V7 对油价按完整日期序列进行插值，并构造 oil_ma7、oil_ma30、oil_ma90、oil_chg7、oil_chg30 等特征，用于刻画短期和中期油价变化。节假日部分区分 national、regional、local holiday，同时加入 is_holiday、is_work_day、days_to_holiday、days_from_holiday 与 holiday_type_enc，捕捉节前节后销量变化。",
    )

    add_heading(doc, "5.3 时间序列特征", 2)
    add_para(
        doc,
        "V7 的核心提升来自时间序列特征。模型使用 log1p(sales) 作为训练目标，并基于 store_nbr-family 组合构造 lag_16、lag_17、lag_18、lag_19、lag_20、lag_21、lag_28、lag_35、lag_56、lag_84、lag_112、lag_182、lag_364 等滞后特征。这些 lag 同时覆盖预测窗口之后可用的最近历史、周周期、月周期、季度周期和年周期。",
    )
    add_para(
        doc,
        "在 lag 基础上，V7 继续构造 rolling mean、rolling std、EWM、momentum 等统计特征，并对促销数量 onpromotion 建立对应的滞后和滚动特征。这样模型不仅看到某一天的历史销量，还能看到最近一段时间的平均水平、波动程度和变化方向。",
    )

    add_heading(doc, "5.4 趋势、Fourier 与 STL 风格分解", 2)
    add_para(
        doc,
        "为了刻画不同时间尺度的周期性，V7 使用 Periodogram 思路选择 Fourier 频率，加入 3.5、7、30、365 等周期的 sin/cos 特征。趋势部分加入 Trend、Trend²、Family×Trend、Cluster×Trend，使模型能够学习不同品类和门店集群的长期增长或下降差异。",
    )
    add_para(
        doc,
        "此外，V7 构造了 STL 风格分解特征：通过历史 lag 和滚动均值近似 trend、seasonal、residual、seasonal strength 等信息。虽然没有对每条序列真正运行完整 STL 分解，但这种近似特征在树模型中能够提供类似“长期水平 + 季节偏离 + 异常残差”的信息。",
    )

    add_heading(doc, "5.5 交互统计与编码特征", 2)
    add_para(
        doc,
        "Store Sales 的数据结构高度分层，因此 V7 构造了多种交互统计特征，包括 store-family、city-family、state-family、cluster-family、store-month、family-dayofweek、family-month、store-dayofweek 等组合统计。V7 还加入 family_rank、store_rank、cluster_rank 等排序编码，用于表达不同类别的相对销售能力。",
    )
    add_para(
        doc,
        "促销与节假日之间也存在交互作用：节假日附近促销可能产生放大效应，周末与节假日叠加时销量行为也会变化。因此 V7 增加 promo_x_holiday、weekend_x_holiday、holiday type one-hot 等增强特征。",
    )

    add_heading(doc, "5.6 模型训练、融合与后处理", 2)
    add_para(
        doc,
        "V7 使用 31 天时间窗口作为验证集，使验证集更接近未来预测任务。模型层面采用 LightGBM、XGBoost 和 CatBoost 的集成：LightGBM 通过 Optuna 搜索参数并训练多个 seed，XGBoost 使用多组深度和学习率配置，CatBoost 作为补充模型。最终在验证集上搜索模型权重，以 log-space 预测进行融合。",
    )
    add_para(
        doc,
        "后处理方面，V7 对闭店或历史长期零销量组合进行置零处理，并尝试季节性预测与模型预测的自适应融合。这个版本最终 Public Score 约为 0.40，是本实验中最重要的自主建模成果，也为后续 V3 的融合优化提供了经验基础。",
    )

    add_heading(doc, "6 V3 融合优化版", 1)
    add_para(
        doc,
        "在 V2 已经形成完整特征工程和集成模型后，继续单纯堆叠复杂模型的收益有限。V3 阶段自然转向融合策略与工程稳定性优化：将不同训练窗口、不同树模型和不同提交候选拆成独立组件，统一输入输出格式，并把它们变成可比较、可验证、可调权重的实验对象。",
    )
    add_para(
        doc,
        "具体工作包括：将实验代码整理为 AutoDL 可运行脚本；修复不同库版本带来的兼容问题；统一 Kaggle/AutoDL 输入输出路径；生成多个候选提交文件；对 LGBM、XGBoost、分层融合等组件进行组合；比较 linear blend 与 log blend；最后用本地权重搜索选择更稳的融合比例。",
    )
    weights = read_csv("best_component_weights.csv")
    add_df_table(doc, weights, "表3 最终四组件权重")
    add_picture(doc, figures["weights"], "图5 四组件融合权重")
    add_para(
        doc,
        "最终采用的本地最优权重为 lgb_full 32.5%、lgb_2015 32.5%、xgb_full 35.0%、xgb_2015 0%。融合方式采用 log-space，即先计算 log1p(pred)，再按权重平均并通过 expm1 还原。相比直接线性融合，log-space 对极端大销量预测更稳，也更符合 RMSLE 指标的形式。",
    )
    add_picture(doc, figures["submission_dist"], "图6 候选提交文件预测分布检查")

    add_heading(doc, "7 实验结果与排名", 1)
    raw_scores = read_csv("score_progression.csv")
    scores = pd.DataFrame(
        {
            "stage": [
                "V2 自研增强模型",
                "V3-1 时序组件实验",
                "V3-2 分层融合实验",
                "V3-3 树模型组件实验",
                "V3-4 log-space 融合",
                "V3-5 加入轻量 LGBM 组件",
                "V3-6 本地权重搜索",
            ],
            "public_score": raw_scores["public_score"].astype(float).round(5),
            "category": [
                "自研特征集成",
                "组件实验",
                "组件实验",
                "组件实验",
                "融合优化",
                "融合优化",
                "权重搜索",
            ],
        }
    )
    add_df_table(doc, scores, "表4 分数迭代记录")
    add_picture(doc, figures["score"], "图7 Public Score 迭代过程", 6.5)
    add_para(
        doc,
        "从成绩变化看，V2 自研增强模型已经比基础版更完整，但单模型/单套特征继续优化的边际收益下降。V3 通过组件融合和权重搜索，将成绩从 0.37946、0.37936 附近继续推进到 0.37927。由于该赛题评分越低越好，0.0001 级别的提升在榜单前段也具有实际意义。",
    )
    add_para(
        doc,
        "Kaggle 页面截图显示，在 Public Score 0.37932 时，账号 Simoni Fretta 排名约第 298 名；之后最终提交提升到 0.37927，最终名次以 Kaggle 页面显示为准。实验过程完整体现了提交、查看结果、分析差距、改进算法并继续提高排名的课程要求。",
    )

    add_heading(doc, "8 总结与反思", 1)
    add_para(
        doc,
        "本次实验从 V1 到 V3 形成了清晰的迭代路线。V1 解决了从数据读取到提交的完整链路；V2 是主要自主工作，围绕时间序列预测构造了大量业务相关特征，并完成多模型融合；V3 则侧重于工程化优化和融合策略，通过组件化输出、log blend 和权重搜索进一步降低 Public RMSLE。",
    )
    add_para(
        doc,
        "实验中也暴露出一些问题。首先，公开榜单分数与本地验证并不完全一致，本地最优权重不一定总能带来榜单最优；其次，特征数量和模型数量增加后，训练成本显著上升，需要在分数提升和时间成本之间权衡；最后，过度依赖 Public Leaderboard 可能产生过拟合风险，后续应构建更稳健的时间序列交叉验证方案。",
    )
    add_para(
        doc,
        "后续改进方向包括：设计多时间窗口验证集，评估不同节假日和促销窗口；细分低销量、闭店和新品类样本；对融合权重进行更稳定的约束搜索；同时保留 V7 的自研特征框架，将 V3 中验证有效的组件化思路逐步沉淀到自己的代码体系中。",
    )

    add_heading(doc, "资料与实现说明", 1)
    add_para(
        doc,
        "实验过程中查阅了 Kaggle 官方数据说明、讨论区经验和公开资料，用于理解赛题特点与校验建模方向。最终实验重点落在自己的实现工作上：V1/V2 完成从基础建模到自研特征集成，V3 完成代码整理、路径适配、运行修复、组件拆分、权重搜索、提交文件分布检查和最终提交验证。",
    )

    doc.save(OUT_DOCX)


def main() -> None:
    ensure_dirs()
    figures = {
        "monthly": line_chart_monthly(),
        "top_family": bar_chart_top_family(),
        "store_type": bar_chart_store_type(),
        "workflow": workflow_chart(),
        "weights": component_weights_chart(),
        "submission_dist": submission_distribution_chart(),
        "score": score_progression_chart(),
    }
    build_markdown()
    build_docx(figures)
    print(f"Saved DOCX: {OUT_DOCX}")
    print(f"Saved MD:   {OUT_MD}")
    print(f"Word figures: {WORD_FIG_DIR}")


if __name__ == "__main__":
    main()
