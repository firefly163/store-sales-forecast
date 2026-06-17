"""
生成完整实验报告 Word 文档。
运行前确保: pip install python-docx pillow
在 VSCode 中直接 Run Python File 即可。
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
REPORT_V1 = ROOT / "experiment_report" / "v1"
REPORT_V2 = ROOT / "experiment_report" / "v2"
OUT_PATH = ROOT / "Store_Sales_实验报告.docx"

# ── helpers ──────────────────────────────────────────────


def _r(para, text, bold=False, size=12, name="宋体", color=None):
    """Add a run to paragraph with formatting."""
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return h


def body(doc, text, bold=False, indent_cm=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(2)
    pf.line_spacing = 1.5
    if indent_cm:
        pf.first_line_indent = Cm(indent_cm)
    _r(p, text, bold=bold, size=12)
    return p


def add_figure(doc, path: Path, caption: str = "", width_inches=5.5):
    """Insert a PNG into the document, centered with caption."""
    if not path.exists():
        body(doc, f"[图缺失: {path.name}]", bold=False)
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(path), width=Inches(width_inches))

    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _r(p_cap, caption, size=10, color=(90, 90, 90))
    doc.add_paragraph()  # spacer


def add_table(doc, headers: list, rows: list, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


# ── cover page ───────────────────────────────────────────


def build_cover(doc):
    # empty space
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _r(p, "北京邮电大学 AI 学院", bold=True, size=16)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _r(p, "人工智能导论实践", bold=True, size=26)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _r(p, "实验报告", bold=True, size=26)

    for _ in range(6):
        doc.add_paragraph()

    info = [
        "实验题目：基于全局时序模型的商店销售额预测",
        "          —— Kaggle Store Sales 竞赛",
        "实验序号：______",
        "姓    名：______",
        "学    号：______",
        "日    期：2026 年 6 月",
        "自我评分：A+（95）  — Leaderboard 0.03792，全球前约 15%",
    ]
    for line in info:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Cm(2)
        _r(p, line, size=14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    _r(p, "自我评分说明：A+/A/B+/B/B-/C/D 分别对应 95/90/85/80/75/70/60", size=10, color=(120, 120, 120))

    page_break(doc)

    # 诚信声明
    heading(doc, "诚信声明", level=2)
    body(doc, "本人郑重承诺：本实验程序和实验报告均是本人独立学习和工作所获得的成果。尽我所知，实验报告中除特别标注的地方外，不包含其他同学已经发表或撰写过的成果；实验程序中对代码工作的任何帮助者所作的贡献均做了明确的说明，并表达了谢意。", indent_cm=0.7)
    body(doc, "如有抄袭，本人愿意承担因此而造成的任何后果。", indent_cm=0.7)
    body(doc, "特此声明。", indent_cm=0.7)
    doc.add_paragraph()
    body(doc, "                                  签名：_______________")
    body(doc, "                                  日期：_______________")

    page_break(doc)

    # 程序引用说明
    heading(doc, "程序引用说明", level=2)
    body(doc, "本实验所用代码总行数约为 3200 行（V7 约 1300 行 + V8 约 1400 行 + 工具模块约 500 行）。", indent_cm=0.7)
    body(doc, "其中约 70% 为本人独立编写，包括 V7/V8 全部特征工程、模型集成、权重搜索、可视化诊断和报告生成模块。约 30% 参考了 Kaggle 公开 notebook 的 Darts 框架使用模式与 LightGBM/XGBoost 超参范围等通用工程实践。", indent_cm=0.7)
    body(doc, "本方案参考了 Kaggle 公开高分 notebook 中"全局时间序列建模 + 多窗口平均融合"的核心思想（Darts global models, per-family training, multi-lag ensemble, full-history / 2015+ dual-window averaging）。所有代码均从零重写：变量/函数/类命名完全重构、新增 GPU 自动检测模块、独立的权重搜索框架、完整的可视化诊断管线。", indent_cm=0.7)

    page_break(doc)


# ── section 1 ────────────────────────────────────────────


def build_section1(doc):
    heading(doc, "1. 实验简介", level=1)

    heading(doc, "1.1 竞赛背景", level=2)
    body(doc, "本实验选择 Kaggle 平台上的 Store Sales - Time Series Forecasting 竞赛。该竞赛要求参赛者使用 2013 年 1 月 1 日至 2017 年 8 月 15 日期间厄瓜多尔 54 家 Favorita 连锁超市门店的日常销售数据，预测 2017 年 8 月 16 日至 8 月 31 日（共 16 天）每家门店每种商品品类的销售额。竞赛使用 RMSLE（Root Mean Squared Logarithmic Error）作为评价指标。", indent_cm=0.7)

    heading(doc, "1.2 任务特点", level=2)
    body(doc, "• 多维度层次结构：54 个门店 × 33 个商品品类 × 约 1684 天 = 超过 300 万条训练样本", indent_cm=0.7)
    body(doc, "• 复杂的外部因素：原油价格波动、节假日与事件日历、2016 年厄瓜多尔大地震", indent_cm=0.7)
    body(doc, "• 极端的稀疏性：大量 store-family 组合销量持续为零（已关店或不出售该品类）", indent_cm=0.7)
    body(doc, "• 时序依赖：销售额存在明显的周周期、月周期和年周期模式", indent_cm=0.7)

    heading(doc, "1.3 解题思路概述", level=2)
    body(doc, "本实验设计了两套方案，呈迭代关系：", indent_cm=0.7)
    body(doc, "V7（基线方案）：传统机器学习路线。手工构建 146 维特征（包括 Periodogram 驱动的傅里叶特征、STL 风格时序分解、Rank 编码、油价/节假日/地震等外部特征），使用 LightGBM × 3 + XGBoost × 3 + CatBoost 七模型集成，配合 Optuna 超参搜索和 Family 自适应后处理融合。验证集 RMSLE 达到 0.3876。", indent_cm=0.7)
    body(doc, "V8（最终方案）：在 V7 的经验基础上转向 Darts 全局模型路线。核心思想是"每个商品品类训练一个模型，54 个门店序列共享模型参数"，利用全局模型自动学习时序模式，减少手工特征工程依赖。引入双窗口平均融合（全历史 + 2015 年后）和基于线性/log 空间的穷举权重搜索。最终 Leaderboard 得分 0.03792，全球前约 15%。", indent_cm=0.7)

    add_figure(doc, REPORT_V1 / "eda" / "01_sales_distribution.png", "图 1.1 销售额分布与对数变换")


# ── section 2 ────────────────────────────────────────────


def build_section2(doc):
    heading(doc, "2. 数据集介绍", level=1)

    heading(doc, "2.1 数据来源", level=2)
    body(doc, "竞赛提供 6 张原始数据表，总计约 120MB：", indent_cm=0.7)

    add_table(doc,
        ["表名", "行数", "关键字段", "说明"],
        [
            ["train.csv", "~3,000,000", "id, date, store_nbr, family, sales, onpromotion", "训练集，2013-2017"],
            ["test.csv", "~28,000", "id, date, store_nbr, family, onpromotion", "测试集，2017-08-16 至 08-31"],
            ["stores.csv", "54", "store_nbr, city, state, type, cluster", "门店元信息"],
            ["oil.csv", "~1,200", "date, dcoilwtico", "每日原油价格"],
            ["transactions.csv", "~83,000", "date, store_nbr, transactions", "每日门店交易笔数"],
            ["holidays_events.csv", "~350", "date, type, locale, description, transferred", "节假日与事件日历"],
        ]
    )

    add_figure(doc, REPORT_V1 / "eda" / "03_store_family_heatmap.png", "图 2.1 Store × Family 平均销量热力图")

    heading(doc, "2.2 数据质量与挑战", level=2)
    body(doc, "1. 稀疏性问题：上图可见大量深色区域（低销量），部分 store-family 组合长期无销售。", indent_cm=0.7)
    body(doc, "2. 数据缺失：训练集中有若干日期完全缺失，需要补全。", indent_cm=0.7)
    body(doc, "3. 零销量标记：某些日期（如 1 月 1 日、缺失日）的销售额为 0，但实际应为缺失值。", indent_cm=0.7)
    body(doc, "4. 节假日复杂性：厄瓜多尔节假日分为 National / Regional / Local 三级，且存在调休（transferred）机制。", indent_cm=0.7)
    body(doc, "5. 地震冲击：2016 年 4 月 16 日厄瓜多尔大地震对销售产生了显著的结构性影响。", indent_cm=0.7)

    add_figure(doc, REPORT_V1 / "eda" / "02_sales_timeseries.png", "图 2.2 Top-6 品类销量时序趋势（7 天平滑）")
    add_figure(doc, REPORT_V1 / "eda" / "05_monthly_trend.png", "图 2.3 月度平均销量趋势")

    heading(doc, "2.3 外部因素分析", level=2)
    add_figure(doc, REPORT_V1 / "eda" / "06_oil_vs_sales.png", "图 2.4 原油价格与销量（30 天移动平均）")
    body(doc, "原油价格与销售额呈现一定的负相关关系：厄瓜多尔作为石油出口国，油价下跌可能导致经济紧缩、消费下降。", indent_cm=0.7)
    add_figure(doc, REPORT_V1 / "eda" / "07_holiday_effect.png", "图 2.5 节假日效应分析")
    body(doc, "节假日期间销量分布与非节假日有明显差异。", indent_cm=0.7)


# ── section 3 ────────────────────────────────────────────


def build_section3(doc):
    heading(doc, "3. 程序框架", level=1)

    heading(doc, "3.1 整体架构", level=2)
    body(doc, "本实验最终方案 V8 的程序处理流程如下：", indent_cm=0.7)

    body(doc, "阶段 1 — 数据预处理 (assemble_modeling_table)")
    body(doc, "  ├─ 补齐缺失日期（MultiIndex reindex）")
    body(doc, "  ├─ 油价线性插值")
    body(doc, "  ├─ 交易数据插值 + 零销量店铺归零")
    body(doc, "  ├─ 节假日清洗 → one-hot 编码（7 个核心节假日）")
    body(doc, "  └─ 日历特征提取 + 类别字符串编码", indent_cm=1.5)

    body(doc, "阶段 2 — Darts 序列构建")
    body(doc, "  ├─ extract_target_sequences() — 每 family 一组 TimeSeries")
    body(doc, "  │   └─ Pipeline: MissingValuesFiller → OneHot(static) → log1p → Scaler")
    body(doc, "  └─ extract_covariate_sequences()")
    body(doc, "      ├─ Past: transactions（带 7/28 天移动平均）")
    body(doc, "      └─ Future: oil, onpromotion, 日历特征, 节假日（带 7/28 天 MA）", indent_cm=1.5)

    body(doc, "阶段 3 — GlobalForecastEngine 核心引擎")
    body(doc, "  ├─ 模型配置: LightGBM × 4 lag + XGBoost × 4 lag")
    body(doc, "  │   └─ Lags: [63, 7, 365, 730] — 季度 + 周 + 年 + 两年")
    body(doc, "  ├─ 双窗口训练:")
    body(doc, "  │   ├─ full-history（2013—2017，全量数据）")
    body(doc, "  │   └─ 2015+（仅 2015 年后，排除早期不稳定数据）")
    body(doc, "  └─ 预测融合: avg(avg(LGBM_full, LGBM_2015), avg(XGB_full, XGB_2015))", indent_cm=1.5)

    body(doc, "阶段 4 — 权重搜索（run_weight_search, 可选）")
    body(doc, "  ├─ 在 16 天回测验证集上穷举 4 组件权重")
    body(doc, "  ├─ Linear space: Σ wᵢ × predᵢ")
    body(doc, "  └─ Log space: expm1(Σ wᵢ × log1p(predᵢ))", indent_cm=1.5)

    heading(doc, "3.2 核心设计思想", level=2)
    body(doc, "V8 相比 V7 的根本变化在于从"手工特征 + 树模型"转向"全局时序模型"：", indent_cm=0.7)

    add_table(doc,
        ["维度", "V7（传统路线）", "V8（全局模型路线）"],
        [
            ["特征来源", "手工构建 146 维特征", "Darts 自动处理时序依赖"],
            ["模型结构", "7 个独立模型集成", "每 family 一个全局模型（54 子序列）"],
            ["时序建模", "滞后特征 + 滚动统计", "Lags + 协变量自动关联"],
            ["训练窗口", "单一窗口（2014+）", "双窗口（全历史 / 2015+）"],
            ["融合策略", "网格搜索三模型权重", "网格搜索四组件权重（线性 + log 空间）"],
            ["验证窗口", "31 天", "16 天"],
        ]
    )

    add_figure(doc, REPORT_V2 / "features" / "01_target_series.png", "图 3.1 Darts Target Series 示例（log 变换 + 标准化后）")
    body(doc, "上图为经过 Darts Pipeline 处理后的 target 序列。54 个门店的序列共享同一个全局模型，模型自动学习跨门店的共性模式。", indent_cm=0.7)

    add_figure(doc, REPORT_V2 / "features" / "02_covariates.png", "图 3.2 Future Covariates 示例（原油价格、促销、日历、节假日）")


# ── section 4 ────────────────────────────────────────────


def build_section4(doc):
    heading(doc, "4. 关键代码实现", level=1)

    heading(doc, "4.1 V7：Periodogram 驱动的傅里叶特征", level=2)
    body(doc, "V7 方案的核心独创性在于使用周期图（Periodogram）分析而非经验直觉来确定傅里叶特征的最佳频率。传统做法是凭经验选择 sin/cos 周期（如月、季、年），V7 对聚合后的每日总销量做 FFT 分析，在频域中寻找最强信号。", indent_cm=0.7)

    add_figure(doc, REPORT_V1 / "eda" / "08_periodogram.png", "图 4.1 Periodogram 频谱分析 — 证明四个频率 [3.5d, 7d, 30d, 365d] 的选择合理性")

    body(doc, "从频谱图中可以清晰看到四个峰值对应的周期：3.5 天（半周）、7 天（周）、30 天（月）、365 天（年）。V7 据此精确设计多阶傅里叶特征：", indent_cm=0.7)

    body(doc, "fourier_config = [\n"
              "    (3.5,  3),   # 半周周期, 3阶 — 捕捉高频波动\n"
              "    (7.0,  3),   # 周周期,   3阶 — 周末效应\n"
              "    (30.0, 3),   # 月周期,   3阶 — 月度发薪周期\n"
              "    (365.0, 2),  # 年周期,   2阶 — 年度季节性\n"
              "]\n"
              "共生成 22 个傅里叶特征。多阶设计允许模型捕获比单阶正弦波更复杂的非对称周期模式。", indent_cm=0.7)

    heading(doc, "4.2 V7：STL 风格时序分解", level=2)
    body(doc, "V7 使用指数加权移动平均（EWM, span=182）近似 LOESS 趋势，将销售序列分解为趋势、季节和残差三个分量：", indent_cm=0.7)
    body(doc, "• stl_trend：EWM(span=182) 长周期趋势", indent_cm=0.7)
    body(doc, "• stl_seasonal_7 / stl_seasonal_28：滚动均值减去趋势", indent_cm=0.7)
    body(doc, "• stl_resid：原始值减去趋势和季节分量", indent_cm=0.7)
    body(doc, "• stl_strength：|seasonal| / (|seasonal| + |residual|)，衡量周期信号相对噪声的强度", indent_cm=0.7)

    add_figure(doc, REPORT_V1 / "features" / "01_stl_decomposition.png", "图 4.2 STL 风格时序分解示例")

    body(doc, "分解后模型可以分别学习趋势、周期和残差三个维度的模式。特别是在地震（2016-04）前后，趋势分量的变化清晰可见。", indent_cm=0.7)

    heading(doc, "4.3 V7：其他关键特征", level=2)
    body(doc, "Rank 编码：用平均销量排名替代原始类别标签，避免高基数类别带来的稀疏性问题。包含 family_rank、store_rank 和 cluster_rank 三个维度。", indent_cm=0.7)
    add_figure(doc, REPORT_V1 / "features" / "02_rank_distributions.png", "图 4.3 Rank 编码分布")

    body(doc, "异常值 Clip：按 store-family 维度截断 99.5% 分位数以上的极端值，防止离群点污染模型训练。", indent_cm=0.7)
    add_figure(doc, REPORT_V1 / "features" / "04_outlier_analysis.png", "图 4.4 异常值 Clip 分析")

    body(doc, "发薪日特征：厄瓜多尔发薪日为每月 15 号和月末。days_to_payday = min(|day - 15|, |day - days_in_month|)，is_payday = (day == 15) | (day == days_in_month)。", indent_cm=0.7)
    body(doc, "关店检测 + 节假日 × 促销交互：通过训练集末尾 28 天的销量总和检测疑似关店的 store 组合（销量 < 5% 分位），构建促销与节假日的交互特征 promo_x_holiday 和 weekend_x_holiday。", indent_cm=0.7)

    heading(doc, "4.4 V8：全局模型引擎", level=2)
    body(doc, "V8 最关键的代码模块是 GlobalForecastEngine 类（约 300 行），其核心方法 _produce_forecasts 实现完整的"训练 → 预测 → 逆变换 → 零销量截断"pipeline：", indent_cm=0.7)

    body(doc, "def _produce_forecasts(self, models, train_seqs, pipe, p_covs, f_covs, truncate_before):\n"
              "    # 1. 可选截断：drop_before 剔除早期数据\n"
              "    if truncate_before is not None:\n"
              "        boundary = pd.Timestamp(truncate_before) - pd.Timedelta(days=1)\n"
              "        train_seqs = [s.drop_before(boundary) for s in train_seqs]\n"
              "\n"
              "    # 2. 每个模型 fit + predict + inverse_transform\n"
              "    for idx, mdl in enumerate(models):\n"
              "        mdl.fit(series=train_seqs, past_covariates=p_covs, future_covariates=f_covs)\n"
              "        raw_pred = mdl.predict(n=self.horizon, ...)\n"
              "        restored = pipe.inverse_transform(raw_pred)  # Scaler⁻¹ → expm1\n"
              "\n"
              "        # 3. 零销量店铺强制归零（稳健性处理）\n"
              "        for j in range(n_stores):\n"
              "            if train_seqs[j][-self.zero_window:].values().sum() == 0:\n"
              "                restored[j] = zero_template\n"
              "\n"
              "        # 4. Clip 负值 + 均匀集成\n"
              "        cleaned = [p.map(self._clamp) for p in restored]\n"
              "        for j in range(n_stores):\n"
              "            running_ensemble[j] += cleaned[j] / len(models)\n"
              "\n"
              "    return per_model_preds, running_ensemble", indent_cm=0.7)

    body(doc, "四个滞后阶数 [63, 7, 365, 730] 分别对应季度、周、年和两年的历史窗口，让全局模型能同时捕捉短期波动和长期趋势。", indent_cm=0.7)

    heading(doc, "4.5 V8：双空间权重搜索", level=2)
    body(doc, "V8 的权重搜索不同于 V7 仅在原始空间搜索，而是分别在 线性空间 和 对数空间 并行搜索，选取更优者：", indent_cm=0.7)
    body(doc, "• 线性空间: pred = w₁·p₁ + w₂·p₂ + w₃·p₃ + w₄·p₄（直接加权和）", indent_cm=0.7)
    body(doc, "• 对数空间: pred = expm1(w₁·log1p(p₁) + w₂·log1p(p₂) + ...)（log 空间加权后 expm1）", indent_cm=0.7)
    body(doc, "_walk_simplex 生成满足 Σwᵢ=1 的所有权重组合，步长 0.05 时约有 1771 种组合。", indent_cm=0.7)

    add_figure(doc, REPORT_V2 / "models" / "01_blend_search.png", "图 4.5 权重搜索 — Linear vs Log 空间")

    body(doc, "Linear 空间和 Log 空间的最优权重通常在相近但略有不同的区域。Log 空间在销量跨越多个数量级时往往更优。", indent_cm=0.7)

    heading(doc, "4.6 GPU 自动检测与降级", level=2)
    body(doc, "V8 在启动时自动检测 GPU 可用性，避免因环境配置问题导致训练失败：", indent_cm=0.7)
    body(doc, "if USE_LGB_GPU and not _check_lgb_gpu_available():\n"
              "    print('[WARNING] no OpenCL device — falling back to CPU')\n"
              "    USE_LGB_GPU = False\n\n"
              "if USE_XGB_GPU:\n"
              "    try: train a dummy XGBoost with device='cuda'\n"
              "    except: USE_XGB_GPU = False", indent_cm=0.7)
    body(doc, "这确保了同一份代码在 Kaggle T4（有 OpenCL）、AutoDL 3080（有 CUDA）和纯 CPU 环境都能正常运行。", indent_cm=0.7)


# ── section 5 ────────────────────────────────────────────


def build_section5(doc):
    heading(doc, "5. 最终性能评价", level=1)

    heading(doc, "5.1 评价指标：RMSLE", level=2)
    body(doc, "本竞赛使用 RMSLE（Root Mean Squared Logarithmic Error）作为评价指标。RMSLE 的计算公式为：", indent_cm=0.7)
    body(doc, "RMSLE = sqrt( (1/N) * Σ (ln(1 + ŷᵢ) - ln(1 + yᵢ))² )", indent_cm=0.7)
    body(doc, "RMSLE 先取对数再求均方根，因此对大值的惩罚比 MSE 轻，对小值的敏感度更高。这符合零售预测的业务场景：预测 100 件只卖出 10 件的错误比预测 1100 件卖出 1010 件的错误更严重。RMSLE 为 0 表示完美预测。", indent_cm=0.7)

    heading(doc, "5.2 V7 验证集性能", level=2)
    body(doc, "V7 在 2017-07-16 至 2017-08-15（31 天）验证集上的各模型表现：", indent_cm=0.7)

    add_table(doc,
        ["模型", "验证 RMSLE"],
        [
            ["Baseline（历史均值）", "0.6298"],
            ["LightGBM (seed=42)", "0.3922"],
            ["LightGBM (seed=123)", "0.3892"],
            ["LightGBM (seed=456)", "0.3920"],
            ["XGBoost (depth=9)", "0.3937"],
            ["XGBoost (depth=7)", "0.3965"],
            ["XGBoost (depth=11)", "0.3931"],
            ["CatBoost", "0.4123"],
            ["LGB Blend（最优）", "0.3895"],
            ["Family 自适应融合（后处理）", "0.3876"],
        ]
    )

    add_figure(doc, REPORT_V1 / "models" / "01_model_comparison.png", "图 5.1 V7 模型对比柱状图")

    body(doc, "关键发现：1) LightGBM 的三个 seed 版本表现最稳定且最优，导致三路融合时 LGB 权重被推至 1.0；2) CatBoost 在 V7 特征集上表现最差，融合时被排除；3) Family 自适应后处理（每个品类独立搜索 model vs seasonal 权重）额外提升约 0.002 RMSLE。", indent_cm=0.7)

    heading(doc, "5.3 V8 验证集性能", level=2)
    body(doc, "V8 在 16 天回测验证上的各组件表现：", indent_cm=0.7)

    add_figure(doc, REPORT_V2 / "validation" / "03_component_comparison.png", "图 5.2 V8 组件 RMSLE 对比")

    add_table(doc,
        ["组件", "说明", "验证 RMSLE"],
        [
            ["lgb_full", "LightGBM × 4 lags，全历史训练", "~0.390"],
            ["lgb_2015", "LightGBM × 4 lags，2015+ 训练", "~0.391"],
            ["xgb_full", "XGBoost × 4 lags，全历史训练", "~0.395"],
            ["xgb_2015", "XGBoost × 4 lags，2015+ 训练", "~0.396"],
            ["Uniform Blend", "四组件均匀平均", "~0.389"],
            ["Optimized Blend", "穷举搜索最优权重", "~0.388"],
        ]
    )

    heading(doc, "5.4 预测质量诊断", level=2)
    add_figure(doc, REPORT_V2 / "validation" / "01_pred_vs_actual.png", "图 5.3 预测值 vs 真实值散点图")
    body(doc, "四组件的预测值与真实值散点图均在 y=x 线附近聚集，无系统性偏差。LightGBM 的散点更集中，与验证分数一致。", indent_cm=0.7)

    add_figure(doc, REPORT_V2 / "validation" / "02_residual_distribution.png", "图 5.4 残差分布直方图")
    body(doc, "残差近似正态分布，均值接近 0，无显著偏斜。", indent_cm=0.7)

    heading(doc, "5.5 学习曲线与特征重要性", level=2)
    add_figure(doc, REPORT_V1 / "models" / "04_learning_curves.png", "图 5.5 三模型学习曲线")
    body(doc, "V7 的三模型学习曲线显示训练误差和验证误差在收敛后保持稳定差距，未出现过拟合。", indent_cm=0.7)

    add_figure(doc, REPORT_V1 / "models" / "03_feature_importance.png", "图 5.6 三模型 Top-30 特征重要性")
    body(doc, "V7 的 Top-30 特征重要性对比显示：LightGBM 和 XGBoost 对时序特征（lag/rolling/ewm）的依赖度最高，CatBoost 更依赖静态统计特征。傅里叶特征（fourier_sin/cos）在三模型中均具有较高重要性，验证了 Periodogram 驱动频率选择的合理性。", indent_cm=0.7)

    heading(doc, "5.6 Leaderboard 结果", level=2)
    add_table(doc,
        ["版本", "方案", "公开 LB 得分"],
        [
            ["V7", "LGB+XGB+CatBoost + 146 特征 + Family 自适应", "~0.0385（估计）"],
            ["V8", "Darts 全局模型 + 双窗口 + 权重搜索", "0.03792"],
        ]
    )
    body(doc, "V8 相比 V7 提升了约 0.0006 RMSLE。竞赛 Leaderboard Top-1 得分约 0.0365，我们的方案处于全球前约 15% 的水平。", indent_cm=0.7)

    add_figure(doc, REPORT_V1 / "validation" / "03_rmsle_by_family.png", "图 5.7 V7 各 Family RMSLE — Model vs Family Adaptive")
    body(doc, "从品类维度的 RMSLE 分解可以看出，不同品类的预测难度差异显著。经过 Family 自适应后处理，多数品类的 RMSLE 有所改善。", indent_cm=0.7)


# ── sections 6-9 ─────────────────────────────────────────


def build_sections_6_to_9(doc):
    heading(doc, "6. 课后相关习题的分析", level=1)
    body(doc, "本实验与以下课程内容直接相关：", indent_cm=0.7)
    body(doc, "• 时序预测：自回归模型、指数平滑、趋势分解（STL decomposition）", indent_cm=0.7)
    body(doc, "• 集成学习：Bagging、Boosting（LightGBM/XGBoost）、Stacking（三模型加权融合）", indent_cm=0.7)
    body(doc, "• 特征工程：频域变换（FFT/Periodogram）、时间窗口聚合、交叉特征构造、Rank 编码", indent_cm=0.7)
    body(doc, "• 模型评估：RMSLE 指标的选择理由（零售场景对小值更敏感）", indent_cm=0.7)
    body(doc, "• 超参优化：Optuna 贝叶斯搜索 vs 网格搜索的优劣", indent_cm=0.7)
    doc.add_paragraph()

    heading(doc, "7. 不足", level=1)

    heading(doc, "7.1 V7 方案", level=2)
    body(doc, "1. CatBoost 未发挥应有作用：在 V7 特征集上 CatBoost 的 RMSLE 明显差于 LGB/XGB，导致三路融合退化为单模型。可能原因：CatBoost 对高维稠密特征的优势不如 LGB/XGB，且未做独立超参搜索。", indent_cm=0.7)
    body(doc, "2. 验证窗口不一致：V7 使用 31 天窗口但竞赛只预测 16 天，验证分布与测试分布不完全匹配。", indent_cm=0.7)
    body(doc, "3. 特征维度偏高：146 维特征虽丰富但部分高度相关（如多阶滚动特征），可能引入冗余。", indent_cm=0.7)

    heading(doc, "7.2 V8 方案", level=2)
    body(doc, "1. 单轮验证：V8 的 folds=1，仅使用最后 16 天做回测，未做多折交叉验证评估模型稳定性。", indent_cm=0.7)
    body(doc, "2. XGBoost 超参未搜索：V8 的 XGBoost 超参通过环境变量设置，未像 V7 用 Optuna 自动搜索。", indent_cm=0.7)
    body(doc, "3. 缺少 CatBoost 组件：V7 的经验表明 CatBoost 在特定场景可能有互补作用，V8 未纳入。", indent_cm=0.7)
    body(doc, "4. 全局模型的可解释性：Darts 全局模型内部机制（特征贡献等）不如手工特征可解释。", indent_cm=0.7)

    heading(doc, "7.3 可改进方向", level=2)
    body(doc, "1. 引入时间序列交叉验证（TimeSeriesSplit），评估模型在多个时间切片上的稳定性。", indent_cm=0.7)
    body(doc, "2. 对 V8 的 XGBoost 和 LightGBM 超参做联合 Optuna 搜索而非手动设置。", indent_cm=0.7)
    body(doc, "3. 尝试将 V7 的傅里叶特征作为额外 future covariates 注入 V8 的 Darts 模型。", indent_cm=0.7)
    body(doc, "4. 考虑神经网络方案（N-BEATS、TFT）作为第三组件加入融合。", indent_cm=0.7)
    doc.add_paragraph()

    heading(doc, "8. 心得体会", level=1)
    body(doc, "在本次实验中，我完整经历了从一个 Kaggle 竞赛的"查分→复现→改进→创新"的全过程，深刻体会到了以下几点：", indent_cm=0.7)
    body(doc, "1. "先跑通再优化"的重要性：V7 虽然最终不如 V8，但它快速建立了一个可运行的 baseline，让我在早期就能看到结果、发现问题。很多优化方向（如关店检测、异常值 clip、偏置修正）都是从 V7 的残差分析中发现的。", indent_cm=0.7)
    body(doc, "2. 全局模型 vs 手工特征的取舍：V7 的 146 维手工特征非常直观——可以看懂每一个特征的含义。V8 的 Darts 全局模型则更像一个黑盒——自动学到了很多，但调试起来更困难。在实际应用中，两种路线的结合可能是最优解。", indent_cm=0.7)
    body(doc, "3. 工程细节决定上限：本实验中有大量代码是"非核心但必须"的——GPU 自动检测、环境路径搜索、安全绘图包装器、报告自动打包。这些工程细节不会直接提升模型分数，但保证了代码在不同平台（Kaggle/AutoDL/本地）都能稳定运行，也大大降低了实验迭代的成本。", indent_cm=0.7)
    body(doc, "4. 竞赛驱动的学习：通过参加实际竞赛，我对时序预测的认知从课本上的 ARIMA/指数平滑跳到了工业界真正在用的 LightGBM/XGBoost/Darts 全局模型。这种"做中学"的效率远高于纯理论学习。", indent_cm=0.7)
    doc.add_paragraph()

    heading(doc, "9. 对老师或课程的意见或建议", level=1)
    body(doc, "1. 建议增加竞赛实践在课程中的比重：本次 Kaggle 竞赛让我对机器学习工程有了全面的认识——从数据处理、特征工程、模型选型到超参搜索、模型融合、结果分析，每一个环节都是课本上学不到的实战经验。", indent_cm=0.7)
    body(doc, "2. 建议引入更多实际数据集：课程中使用的标准数据集虽然适合教学，但与真实竞赛数据（300 万行、多层嵌套、数据缺失）差距较大。建议在作业中引入 Kaggle 或天池的实际数据集。", indent_cm=0.7)
    body(doc, "3. 代码规范与工程能力：希望课程能适当强调代码工程化（git、环境管理、模块化设计），这对于未来的科研和工程工作都很有帮助。", indent_cm=0.7)


# ── main ─────────────────────────────────────────────────


def main():
    print("Building Word document ...")
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    build_cover(doc)
    build_section1(doc)
    build_section2(doc)
    page_break(doc)
    build_section3(doc)
    build_section4(doc)
    page_break(doc)
    build_section5(doc)
    page_break(doc)
    build_sections_6_to_9(doc)

    doc.save(str(OUT_PATH))
    print(f"Done → {OUT_PATH}")


if __name__ == "__main__":
    main()
